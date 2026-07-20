"""
Generate Thermometer Shift Register Block Documentation (.docx)

Creates a full specification/explainer document for the hyp_adc_cdr_dac_sr_unit
block with diagrams, truth tables, valid/invalid input tables, and reset rules.

Usage:
    python generate_sr_doc.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import io
from datetime import date

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np


# ═══════════════════════════════════════════════════════════════════════════
# Simulation model (from therm_sr_sim_v1_2 — netlist-verified)
# ═══════════════════════════════════════════════════════════════════════════

MAX_ITER = 50

def mux_wr2(in0, in1, sel, sel_rst, in_rst):
    if sel_rst:
        return 1 - in_rst
    return in1 if sel else in0

def evaluate(state, in_pre, in_post, vsel1, vsel0, sel_rst, in_rst):
    s = list(state)
    for it in range(1, MAX_ITER + 1):
        n = [0] * 4
        n[0] = mux_wr2(in_pre,  s[1], vsel1, sel_rst, in_rst[0])
        n[1] = mux_wr2(s[0],   s[2], vsel0, sel_rst, in_rst[1])
        n[2] = mux_wr2(s[3],   s[1], vsel1, sel_rst, in_rst[2])
        n[3] = mux_wr2(in_post, s[2], vsel0, sel_rst, in_rst[3])
        if n == s:
            return n, True, it
        s = n
    return s, False, MAX_ITER


# ═══════════════════════════════════════════════════════════════════════════
# Waveform generation
# ═══════════════════════════════════════════════════════════════════════════

def _digital_wave(ax, steps, values, label, color, yoff=0, height=0.7):
    """Draw a digital waveform with step transitions."""
    for i in range(len(steps) - 1):
        ax.fill_between([steps[i], steps[i + 1]], yoff, yoff + height * values[i],
                        step="post", color=color, alpha=0.35)
        ax.step([steps[i], steps[i + 1]], [yoff + height * values[i]] * 2,
                where="post", color=color, linewidth=1.5)
    ax.text(steps[0] - 0.4, yoff + height * 0.35, label, ha="right", va="center",
            fontsize=7, fontweight="bold")


def generate_sweep_waveform():
    """Simulate full sweep 0→4→0 and produce a waveform image."""
    cw_order = [(0, 1), (0, 0), (1, 0), (1, 1)]
    ccw_order = [(1, 0), (0, 0), (0, 1), (1, 1)]

    state = [0, 0, 0, 0]
    vsel = (1, 1)

    rec = {"step": [0], "code": [0],
           "vsel1": [1], "vsel0": [1],
           "v0": [0], "v1": [0], "v2": [0], "v3": [0]}

    # UP
    t = 0
    for i in range(8):
        nv = cw_order[i % 4]
        state, _, _ = evaluate(state, 1, 0, nv[0], nv[1], 0, [0] * 4)
        t += 1
        rec["step"].append(t)
        rec["code"].append(sum(state))
        rec["vsel1"].append(nv[0]); rec["vsel0"].append(nv[1])
        rec["v0"].append(state[0]); rec["v1"].append(state[1])
        rec["v2"].append(state[2]); rec["v3"].append(state[3])
        vsel = nv
        if sum(state) == 4 and i > 0:
            break

    # DOWN from saturated state
    for i in range(8):
        nv = ccw_order[i % 4]
        state, _, _ = evaluate(state, 1, 0, nv[0], nv[1], 0, [0] * 4)
        t += 1
        rec["step"].append(t)
        rec["code"].append(sum(state))
        rec["vsel1"].append(nv[0]); rec["vsel0"].append(nv[1])
        rec["v0"].append(state[0]); rec["v1"].append(state[1])
        rec["v2"].append(state[2]); rec["v3"].append(state[3])
        if sum(state) == 0 and i > 0:
            break

    N = len(rec["step"])
    steps = list(range(N + 1))

    fig, axes = plt.subplots(4, 1, figsize=(7, 5.5), sharex=True,
                              gridspec_kw={"height_ratios": [1, 1, 2, 1.5]})

    # vsel<1>
    _digital_wave(axes[0], steps, rec["vsel1"], "vsel<1>", "#2196F3")
    axes[0].set_ylim(-0.1, 1.1); axes[0].set_yticks([0, 1])
    axes[0].set_ylabel(""); axes[0].grid(axis="x", alpha=0.3)

    # vsel<0>
    _digital_wave(axes[1], steps, rec["vsel0"], "vsel<0>", "#4CAF50")
    axes[1].set_ylim(-0.1, 1.1); axes[1].set_yticks([0, 1])
    axes[1].set_ylabel(""); axes[1].grid(axis="x", alpha=0.3)

    # vout<3:0>
    colors_v = ["#E91E63", "#FF9800", "#9C27B0", "#607D8B"]
    labels_v = ["vout0", "vout1", "vout2", "vout3"]
    for idx, key in enumerate(["v0", "v1", "v2", "v3"]):
        off = (3 - idx) * 1.15
        _digital_wave(axes[2], steps, rec[key], labels_v[idx], colors_v[idx],
                      yoff=off, height=0.9)
    axes[2].set_ylim(-0.2, 4.8); axes[2].set_yticks([])
    axes[2].grid(axis="x", alpha=0.3)

    # Code
    axes[3].step(steps[:-1], rec["code"], where="post", color="#333", linewidth=2)
    for i, c in enumerate(rec["code"]):
        axes[3].text(i + 0.5, c + 0.15, str(c), ha="center", fontsize=7)
    axes[3].set_ylim(-0.3, 4.8); axes[3].set_yticks(range(5))
    axes[3].set_ylabel("Code", fontsize=8); axes[3].set_xlabel("Step", fontsize=8)
    axes[3].grid(axis="both", alpha=0.3)

    # Annotate UP / DOWN regions
    mid_up = len([c for c in rec["code"] if c <= 4]) // 2
    axes[0].set_title("Full Sweep: UP (CW gray) then DOWN (CCW gray)",
                      fontsize=10, fontweight="bold")

    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=180)
    plt.close(fig)
    buf.seek(0)
    return buf


def generate_reset_waveform():
    """Simulate a reset-to-code-0 (unit OFF) with vsel=11, no toggling."""
    # Unit resets to code 0: in_rst=1111 → vout=0000
    in_rst = [1, 1, 1, 1]
    state = [0, 1, 0, 1]  # garbage

    rec = {"step": [], "sel_rst": [], "vsel1": [], "vsel0": [],
           "v0": [], "v1": [], "v2": [], "v3": [],
           "rst0": [], "rst1": [], "rst2": [], "rst3": []}

    def snap(t, sr, v1, v0):
        rec["step"].append(t)
        rec["sel_rst"].append(sr)
        rec["vsel1"].append(v1); rec["vsel0"].append(v0)
        rec["v0"].append(state[0]); rec["v1"].append(state[1])
        rec["v2"].append(state[2]); rec["v3"].append(state[3])
        rec["rst0"].append(in_rst[0]); rec["rst1"].append(in_rst[1])
        rec["rst2"].append(in_rst[2]); rec["rst3"].append(in_rst[3])

    # t=0: garbage state, sel_rst=0, vsel=11
    snap(0, 0, 1, 1)

    # t=1: set vsel=11, in_rst=1111 (still before reset)
    state_pre = list(state)
    snap(1, 0, 1, 1)

    # t=2: assert sel_rst=1 at vsel=11
    state, _, _ = evaluate(state, 1, 0, 1, 1, 1, in_rst)
    snap(2, 1, 1, 1)

    # t=3: hold under reset (no toggling)
    state, _, _ = evaluate(state, 1, 0, 1, 1, 1, in_rst)
    snap(3, 1, 1, 1)

    # t=4: release reset (sel_rst → 0)
    state, _, _ = evaluate(state, 1, 0, 1, 1, 0, in_rst)
    snap(4, 0, 1, 1)

    # t=5: hold — verify stable
    state, _, _ = evaluate(state, 1, 0, 1, 1, 0, in_rst)
    snap(5, 0, 1, 1)

    N = len(rec["step"])
    steps = list(range(N + 1))

    fig, axes = plt.subplots(4, 1, figsize=(7, 5), sharex=True,
                              gridspec_kw={"height_ratios": [1, 1, 1, 2]})

    # sel_rst
    _digital_wave(axes[0], steps, rec["sel_rst"], "sel_rst", "#F44336")
    axes[0].set_ylim(-0.1, 1.1); axes[0].set_yticks([0, 1])
    axes[0].set_title("Reset Sequence: Unit OFF (in_rst=1111, vsel=11, no toggling)",
                      fontsize=10, fontweight="bold")
    axes[0].grid(axis="x", alpha=0.3)

    # vsel
    _digital_wave(axes[1], steps, rec["vsel1"], "vsel<1>", "#2196F3")
    _digital_wave(axes[1], steps, rec["vsel0"], "vsel<0>", "#4CAF50", yoff=1.3)
    axes[1].set_ylim(-0.1, 2.5); axes[1].set_yticks([])
    axes[1].grid(axis="x", alpha=0.3)

    # in_rst (constant during reset)
    colors_r = ["#90A4AE"] * 4
    for idx, key in enumerate(["rst0", "rst1", "rst2", "rst3"]):
        off = (3 - idx) * 1.15
        _digital_wave(axes[2], steps, rec[key],
                      f"in_rst<{idx}>", colors_r[idx], yoff=off, height=0.9)
    axes[2].set_ylim(-0.2, 4.8); axes[2].set_yticks([])
    axes[2].grid(axis="x", alpha=0.3)

    # vout
    colors_v = ["#E91E63", "#FF9800", "#9C27B0", "#607D8B"]
    for idx, key in enumerate(["v0", "v1", "v2", "v3"]):
        off = (3 - idx) * 1.15
        _digital_wave(axes[3], steps, rec[key],
                      f"vout<{idx}>", colors_v[idx], yoff=off, height=0.9)
    axes[3].set_ylim(-0.2, 4.8); axes[3].set_yticks([])
    axes[3].set_xlabel("Step", fontsize=8)
    axes[3].grid(axis="x", alpha=0.3)

    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=180)
    plt.close(fig)
    buf.seek(0)
    return buf


def generate_zigzag_waveform():
    """Simulate zigzag (UP 3, DOWN 2, UP 1) and produce a waveform image."""
    state = [0, 0, 0, 0]

    cw_order  = [(0, 1), (0, 0), (1, 0), (1, 1)]
    ccw_order = [(1, 0), (0, 0), (0, 1), (1, 1)]

    rec = {"step": [0], "code": [0],
           "vsel1": [1], "vsel0": [1],
           "v0": [0], "v1": [0], "v2": [0], "v3": [0],
           "direction": ["—"]}

    cw_idx = 0
    ccw_idx = 0

    def step_dir(direction, n):
        nonlocal state, cw_idx, ccw_idx
        order = cw_order if direction == "UP" else ccw_order
        idx_ref = cw_idx if direction == "UP" else ccw_idx
        for _ in range(n):
            nv = order[idx_ref % 4]
            state, _, _ = evaluate(state, 1, 0, nv[0], nv[1], 0, [0] * 4)
            t = len(rec["step"])
            rec["step"].append(t)
            rec["code"].append(sum(state))
            rec["vsel1"].append(nv[0]); rec["vsel0"].append(nv[1])
            rec["v0"].append(state[0]); rec["v1"].append(state[1])
            rec["v2"].append(state[2]); rec["v3"].append(state[3])
            rec["direction"].append(direction)
            idx_ref += 1
        if direction == "UP":
            cw_idx = idx_ref
            ccw_map = {(1,1):0, (1,0):1, (0,0):2, (0,1):3}
            ccw_idx = ccw_map.get(nv, 0)
        else:
            ccw_idx = idx_ref
            cw_map = {(1,1):0, (0,1):1, (0,0):2, (1,0):3}
            cw_idx = cw_map.get(nv, 0)

    step_dir("UP", 3)
    step_dir("DN", 2)
    step_dir("UP", 1)

    N = len(rec["step"])
    steps = list(range(N + 1))

    fig, ax = plt.subplots(figsize=(6, 2.5))
    ax.step(steps[:-1], rec["code"], where="post", color="#333", linewidth=2)
    for i, c in enumerate(rec["code"]):
        col = "#2196F3" if rec["direction"][i] == "UP" else (
              "#F44336" if rec["direction"][i] == "DN" else "#999")
        ax.text(i + 0.5, c + 0.18, str(c), ha="center", fontsize=8, color=col)
    ax.set_ylim(-0.3, 4.3); ax.set_yticks(range(5))
    ax.set_ylabel("Code", fontsize=9); ax.set_xlabel("Step", fontsize=9)
    ax.set_title("Zigzag: UP 3, DOWN 2, UP 1  →  expected code = 2",
                 fontsize=10, fontweight="bold")
    ax.grid(alpha=0.3)
    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=180)
    plt.close(fig)
    buf.seek(0)
    return buf


def set_cell(cell, text, bold=False, size=9, align=None, bg=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if bg:
        tc_pr = cell._element.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), bg)
        shd.set(qn("w:val"), "clear")
        tc_pr.append(shd)


def add_table(doc, headers, rows, col_widths=None, header_bg="D9E2F3"):
    n_cols = len(headers)
    n_rows = 1 + len(rows)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        set_cell(table.rows[0].cells[i], h, bold=True, size=9, bg=header_bg)
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            set_cell(table.rows[r_idx + 1].cells[c_idx], str(val), size=9)
    doc.add_paragraph()
    return table


def add_mono(doc, text, size=8):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(size)
    return p


def generate_gray_code_diagram():
    """Generate CW/CCW gray code state diagram showing code direction."""
    fig, ax = plt.subplots(1, 1, figsize=(5.5, 5.5))
    ax.set_xlim(-2.5, 2.5)
    ax.set_ylim(-2.5, 2.5)
    ax.set_aspect('equal')
    ax.axis('off')

    # State positions (circle): 00=top, 01=right, 11=bottom, 10=left
    states = {
        '00': (0, 1.5),
        '01': (1.5, 0),
        '11': (0, -1.5),
        '10': (-1.5, 0),
    }

    # Draw state circles
    for label, (x, y) in states.items():
        circle = plt.Circle((x, y), 0.35, fill=False, linewidth=2.0,
                           edgecolor='#444444')
        ax.add_patch(circle)
        ax.text(x, y, label, ha='center', va='center', fontsize=14,
                fontweight='bold', fontfamily='monospace')

    # CW arrows (blue): 11→01→00→10→11 = UP (+1) — outer ring
    cw_transitions = [('11', '01'), ('01', '00'), ('00', '10'), ('10', '11')]
    # CCW arrows (red): 11→10→00→01→11 = DOWN (-1) — inner ring
    ccw_transitions = [('11', '10'), ('10', '00'), ('00', '01'), ('01', '11')]

    def draw_arrow(src, dst, color, rad):
        """Draw curved arrow between two states."""
        x0, y0 = states[src]
        x1, y1 = states[dst]
        # Shorten arrows to not overlap circles
        dx, dy = x1 - x0, y1 - y0
        length = (dx**2 + dy**2)**0.5
        shrink = 0.40 / length
        sx = x0 + dx * shrink
        sy = y0 + dy * shrink
        ex = x1 - dx * shrink
        ey = y1 - dy * shrink
        ax.annotate('', xy=(ex, ey), xytext=(sx, sy),
                    arrowprops=dict(arrowstyle='->', color=color, lw=2.8,
                                    connectionstyle=f'arc3,rad={rad}'))

    for src, dst in cw_transitions:
        draw_arrow(src, dst, '#1565C0', 0.35)   # blue, curve outward

    for src, dst in ccw_transitions:
        draw_arrow(src, dst, '#C62828', 0.35)    # red, curve outward (opposite direction)

    # Legend
    from matplotlib.lines import Line2D
    legend_elements = [
        Line2D([0], [0], color='#1565C0', lw=2.8, label='CW = UP (+1 code)'),
        Line2D([0], [0], color='#C62828', lw=2.8, label='CCW = DOWN (\u22121 code)'),
    ]
    ax.legend(handles=legend_elements, loc='lower center', fontsize=11,
              framealpha=0.9, ncol=2, bbox_to_anchor=(0.5, -0.08))

    ax.set_title("Gray Code State Transitions (vsel<1:0>)", fontsize=12,
                 fontweight='bold', pad=12)

    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=180, bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)
    return buf


def generate_pair_diagrams():
    """Generate a 2x2 grid showing pair formation for each vsel state."""
    fig, axes = plt.subplots(2, 2, figsize=(9, 7))
    fig.suptitle("Cross-Coupled Pair Formation per vsel State", fontsize=12,
                 fontweight='bold', y=0.98)

    # Physical order left to right: mux3, mux2, mux1, mux0
    # display_order maps logical mux index to physical position (0=left, 3=right)
    display_order = [3, 2, 1, 0]  # mux3 at pos 0, mux0 at pos 3

    configs = [
        {"vsel": "00", "title": "vsel=00: NO PAIR (anchor)",
         "pair": None, "forced": [0, 1, 2, 3],
         "values": {"mux3": "0 (in_post)", "mux2": "0 (=vout3)",
                    "mux1": "1 (=vout0)", "mux0": "1 (in_pre)"},
         "codes": "Code 2 only: 1100 (FORCED)"},
        {"vsel": "10", "title": "vsel=10: PAIR at (mux0 ↔ mux1)",
         "pair": (0, 1), "forced": [3], "followers": [2],
         "values": {"mux3": "0 (in_post)", "mux2": "=pair",
                    "mux1": "PAIR", "mux0": "PAIR"},
         "codes": "pair=0 → 0000 (code 0)\npair=1 → 1110 (code 3)"},
        {"vsel": "11", "title": "vsel=11: PAIR at (mux1 ↔ mux2)",
         "pair": (1, 2), "forced": [], "followers": [0, 3],
         "values": {"mux3": "=pair", "mux2": "PAIR",
                    "mux1": "PAIR", "mux0": "=pair"},
         "codes": "pair=0 → 0000 (code 0)\npair=1 → 1111 (code 4)"},
        {"vsel": "01", "title": "vsel=01: PAIR at (mux2 ↔ mux3)",
         "pair": (2, 3), "forced": [0], "followers": [1],
         "values": {"mux3": "PAIR", "mux2": "PAIR",
                    "mux1": "=pair", "mux0": "1 (in_pre)"},
         "codes": "pair=0 → 1000 (code 1)\npair=1 → 1111 (code 4)"},
    ]

    for idx, cfg in enumerate(configs):
        ax = axes[idx // 2][idx % 2]
        ax.set_xlim(-0.5, 10)
        ax.set_ylim(-2.5, 3.5)
        ax.set_aspect('equal')
        ax.axis('off')
        ax.set_title(cfg["title"], fontsize=10, fontweight='bold', pad=5)

        # Draw 4 mux boxes in physical order: mux3, mux2, mux1, mux0
        box_w, box_h = 1.8, 1.6
        positions = [(0.5, 1), (3, 1), (5.5, 1), (8, 1)]
        phys_names = ["mux3", "mux2", "mux1", "mux0"]
        phys_values = [cfg["values"]["mux3"], cfg["values"]["mux2"],
                       cfg["values"]["mux1"], cfg["values"]["mux0"]]
        # Map logical mux index to physical position for pair coloring
        logical_to_phys = {3: 0, 2: 1, 1: 2, 0: 3}

        for i, (cx, cy) in enumerate(positions):
            logical_idx = display_order[i]  # which logical mux is at this position

            # Color based on role
            if cfg["pair"] and logical_idx in cfg["pair"]:
                fc = '#FFCDD2'  # red tint for pair
                ec = '#C62828'
                lw = 2.5
            elif logical_idx in cfg.get("forced", []):
                fc = '#BBDEFB'  # blue tint for forced
                ec = '#1565C0'
                lw = 1.5
            else:
                fc = '#E8F5E9'  # green tint for followers
                ec = '#388E3C'
                lw = 1.5

            rect = mpatches.FancyBboxPatch(
                (cx - box_w/2, cy - box_h/2), box_w, box_h,
                boxstyle="round,pad=0.05", edgecolor=ec,
                facecolor=fc, linewidth=lw)
            ax.add_patch(rect)
            ax.text(cx, cy + 0.25, phys_names[i], ha='center', va='center',
                    fontsize=8, fontweight='bold')
            ax.text(cx, cy - 0.35, phys_values[i], ha='center', va='center',
                    fontsize=6.5, color='#333')

        # Draw pair arrows if pair exists
        if cfg["pair"]:
            p0_log, p1_log = cfg["pair"]
            p0_phys = logical_to_phys[p0_log]
            p1_phys = logical_to_phys[p1_log]
            # Ensure p0_phys < p1_phys for consistent drawing
            if p0_phys > p1_phys:
                p0_phys, p1_phys = p1_phys, p0_phys
            x0, y0 = positions[p0_phys]
            x1, y1 = positions[p1_phys]
            # Top arrow (left to right)
            ax.annotate('', xy=(x1 - box_w/2, y1 + 0.5),
                        xytext=(x0 + box_w/2, y0 + 0.5),
                        arrowprops=dict(arrowstyle='->', color='#C62828',
                                        lw=2, connectionstyle='arc3,rad=-0.3'))
            # Bottom arrow (right to left)
            ax.annotate('', xy=(x0 + box_w/2, y0 - 0.5),
                        xytext=(x1 - box_w/2, y1 - 0.5),
                        arrowprops=dict(arrowstyle='->', color='#C62828',
                                        lw=2, connectionstyle='arc3,rad=-0.3'))
            # LATCH label
            mid_x = (x0 + x1) / 2
            ax.text(mid_x, y0 + 1.3, "LATCH", ha='center', va='center',
                    fontsize=8, fontweight='bold', color='#C62828',
                    bbox=dict(boxstyle='round,pad=0.2', facecolor='#FFEBEE',
                              edgecolor='#C62828', linewidth=0.5))

        # Holdable codes text below
        ax.text(5, -1.5, cfg["codes"], ha='center', va='center',
                fontsize=8, fontfamily='monospace',
                bbox=dict(boxstyle='round,pad=0.3', facecolor='#FFF9C4',
                          edgecolor='#F9A825', linewidth=1))

    # Legend
    from matplotlib.lines import Line2D
    legend_ax = fig.add_axes([0.15, 0.01, 0.7, 0.03])
    legend_ax.axis('off')
    legend_elements = [
        mpatches.Patch(facecolor='#FFCDD2', edgecolor='#C62828', lw=1.5,
                       label='Pair (latch)'),
        mpatches.Patch(facecolor='#BBDEFB', edgecolor='#1565C0', lw=1.5,
                       label='Forced by boundary'),
        mpatches.Patch(facecolor='#E8F5E9', edgecolor='#388E3C', lw=1.5,
                       label='Follower (=pair value)'),
    ]
    legend_ax.legend(handles=legend_elements, loc='center', ncol=3, fontsize=9,
                     frameon=False)

    plt.subplots_adjust(left=0.02, right=0.98, top=0.92, bottom=0.08,
                        wspace=0.1, hspace=0.4)
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=180, bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)
    return buf


def generate_block_diagram():
    """Generate a proper block diagram of the SR unit architecture.
    Physical order (matching schematic): mux3, mux2, mux1, mux0 left to right.
    """
    fig, ax = plt.subplots(1, 1, figsize=(9, 5.5))
    ax.set_xlim(-2.5, 19)
    ax.set_ylim(-5.5, 5.0)
    ax.set_aspect('equal')
    ax.axis('off')

    # Physical positions left-to-right: mux3, mux2, mux1, mux0
    mux_positions = [(2, 0), (6.5, 0), (11, 0), (15.5, 0)]
    box_w, box_h = 2.2, 3.0
    mux_names = ["mux3", "mux2", "mux1", "mux0"]
    sel_labels = ["vsel<0>", "vsel<1>", "vsel<0>", "vsel<1>"]
    out_labels = ["vout3", "vout2", "vout1", "vout0"]

    for i, (cx, cy) in enumerate(mux_positions):
        # Draw box
        rect = mpatches.FancyBboxPatch(
            (cx - box_w/2, cy - box_h/2), box_w, box_h,
            boxstyle="round,pad=0.05", edgecolor='black',
            facecolor='#E3F2FD', linewidth=1.5)
        ax.add_patch(rect)

        # Mux name
        ax.text(cx, cy, mux_names[i], ha='center', va='center',
                fontsize=11, fontweight='bold')

        # Port labels inside box
        ax.text(cx - box_w/2 + 0.15, cy + 0.8, "in0", ha='left', va='center',
                fontsize=8, color='#1565C0', fontweight='bold')
        ax.text(cx - box_w/2 + 0.15, cy - 0.8, "in1", ha='left', va='center',
                fontsize=8, color='#C62828', fontweight='bold')

        # sel label on top
        ax.text(cx, cy + box_h/2 + 0.25, sel_labels[i], ha='center',
                va='bottom', fontsize=9, color='#4527A0', fontweight='bold')

        # Output label below
        ax.text(cx, cy - box_h/2 - 0.25, out_labels[i], ha='center',
                va='top', fontsize=9, fontweight='bold', color='#333')

    # Position references:
    # pos 0 = mux3 (x=2), pos 1 = mux2 (x=6.5), pos 2 = mux1 (x=11), pos 3 = mux0 (x=15.5)

    # ── in0 chain arrows (blue, solid) ──
    # in_post → mux3.in0 (from left)
    ax.annotate('', xy=(2 - box_w/2, 0.8), xytext=(-1.5, 0.8),
                arrowprops=dict(arrowstyle='->', color='#1565C0', lw=2))
    ax.text(-1.7, 0.8, "in_post", ha='right', va='center', fontsize=10,
            fontweight='bold', color='#1565C0')

    # vout3 → mux2.in0 (mux3 output at pos0 → mux2 input at pos1, ADJACENT)
    ax.plot([2, 2], [-box_h/2, -box_h/2 - 0.5], color='#1565C0', lw=1.5)
    ax.plot([2, 5.0], [-box_h/2 - 0.5, -box_h/2 - 0.5], color='#1565C0', lw=1.5)
    ax.plot([5.0, 5.0], [-box_h/2 - 0.5, 0.8], color='#1565C0', lw=1.5)
    ax.annotate('', xy=(6.5 - box_w/2, 0.8), xytext=(5.0, 0.8),
                arrowprops=dict(arrowstyle='->', color='#1565C0', lw=1.5))

    # vout0 → mux1.in0 (mux0 output at pos3 → mux1 input at pos2, ADJACENT)
    ax.plot([15.5, 15.5], [-box_h/2, -box_h/2 - 0.5], color='#1565C0', lw=1.5)
    ax.plot([15.5, 9.5], [-box_h/2 - 0.5, -box_h/2 - 0.5], color='#1565C0', lw=1.5)
    ax.plot([9.5, 9.5], [-box_h/2 - 0.5, 0.8], color='#1565C0', lw=1.5)
    ax.annotate('', xy=(11 - box_w/2, 0.8), xytext=(9.5, 0.8),
                arrowprops=dict(arrowstyle='->', color='#1565C0', lw=1.5))

    # in_pre → mux0.in0 (from right)
    ax.annotate('', xy=(15.5 + box_w/2, 0.8), xytext=(17.5, 0.8),
                arrowprops=dict(arrowstyle='->', color='#1565C0', lw=2))
    ax.text(17.7, 0.8, "in_pre", ha='left', va='center', fontsize=10,
            fontweight='bold', color='#1565C0')

    # ── in1 cross-links (red) ──
    # vout1 → mux0.in1 (mux1 at pos2 → mux0 at pos3, adjacent)
    # vout1 → mux2.in1 (mux1 at pos2 → mux2 at pos1, adjacent)
    ax.plot([11, 11], [-box_h/2, -box_h/2 - 2.0], color='#C62828', lw=1.5)
    # Right to mux0
    ax.plot([11, 14.1], [-box_h/2 - 2.0, -box_h/2 - 2.0], color='#C62828', lw=1.5)
    ax.plot([14.1, 14.1], [-box_h/2 - 2.0, -0.8], color='#C62828', lw=1.5)
    ax.annotate('', xy=(15.5 - box_w/2, -0.8), xytext=(14.1, -0.8),
                arrowprops=dict(arrowstyle='->', color='#C62828', lw=1.5))
    # Left to mux2
    ax.plot([11, 5.1], [-box_h/2 - 2.0, -box_h/2 - 2.0], color='#C62828', lw=1.5)
    ax.plot([5.1, 5.1], [-box_h/2 - 2.0, -0.8], color='#C62828', lw=1.5)
    ax.annotate('', xy=(6.5 - box_w/2, -0.8), xytext=(5.1, -0.8),
                arrowprops=dict(arrowstyle='->', color='#C62828', lw=1.5))
    # dot at junction
    ax.plot(11, -box_h/2 - 2.0, 'o', color='#C62828', markersize=5)

    # vout2 → mux1.in1 (mux2 at pos1 → mux1 at pos2, adjacent)
    # vout2 → mux3.in1 (mux2 at pos1 → mux3 at pos0, adjacent)
    ax.plot([6.5, 6.5], [-box_h/2, -box_h/2 - 3.2], color='#C62828', lw=1.5,
            linestyle='--')
    # Right to mux1
    ax.plot([6.5, 9.6], [-box_h/2 - 3.2, -box_h/2 - 3.2], color='#C62828',
            lw=1.5, linestyle='--')
    ax.plot([9.6, 9.6], [-box_h/2 - 3.2, -0.8], color='#C62828', lw=1.5,
            linestyle='--')
    ax.annotate('', xy=(11 - box_w/2, -0.8), xytext=(9.6, -0.8),
                arrowprops=dict(arrowstyle='->', color='#C62828', lw=1.5))
    # Left to mux3
    ax.plot([6.5, 0.6], [-box_h/2 - 3.2, -box_h/2 - 3.2], color='#C62828',
            lw=1.5, linestyle='--')
    ax.plot([0.6, 0.6], [-box_h/2 - 3.2, -0.8], color='#C62828', lw=1.5,
            linestyle='--')
    ax.annotate('', xy=(2 - box_w/2, -0.8), xytext=(0.6, -0.8),
                arrowprops=dict(arrowstyle='->', color='#C62828', lw=1.5))
    # dot at junction
    ax.plot(6.5, -box_h/2 - 3.2, 'o', color='#C62828', markersize=5)

    # ── Legend ──
    from matplotlib.lines import Line2D
    legend_elements = [
        Line2D([0], [0], color='#1565C0', lw=2, label='in0 chain (sel=0 path)'),
        Line2D([0], [0], color='#C62828', lw=2, label='in1: vout1 → mux0, mux2'),
        Line2D([0], [0], color='#C62828', lw=2, linestyle='--',
               label='in1: vout2 → mux1, mux3'),
    ]
    ax.legend(handles=legend_elements, loc='upper right', fontsize=8,
              framealpha=0.9)

    ax.set_title("SR Unit Block Diagram — Physical Layout (mux3..mux0, left to right)",
                 fontsize=11, fontweight='bold', pad=10)

    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=180, bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)
    return buf


def generate():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # ── Title ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Thermometer Shift Register Unit\nhyp_adc_cdr_dac_sr_unit")
    run.font.size = Pt(20)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Block Description & Usage Guide")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph()
    info_table = doc.add_table(rows=3, cols=2)
    info_table.style = "Table Grid"
    info_data = [
        ("Date", date.today().strftime("%Y-%m-%d")),
        ("Version", "1.0"),
        ("Status", "Verified against Spectre netlist"),
    ]
    for i, (k, v) in enumerate(info_data):
        set_cell(info_table.rows[i].cells[0], k, bold=True, size=10)
        set_cell(info_table.rows[i].cells[1], v, size=10)
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 1. OVERVIEW
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("1. Overview", level=1)
    doc.add_paragraph(
        "The hyp_adc_cdr_dac_sr_unit is a 4-mux bidirectional thermometer shift "
        "register that converts gray-coded control signals into a 4-bit thermometer "
        "output (codes 0 through 4). It is the fundamental building block of the "
        "CDR DAC's digital control path."
    )
    doc.add_paragraph(
        "Each gray code step (single-bit toggle of vsel<1:0>) increments or "
        "decrements the thermometer code by exactly 1. The block uses cross-coupled "
        "mux pairs for state retention — no clock or flip-flops required."
    )

    # ── Port table ──
    doc.add_heading("1.1 Port List", level=2)
    add_table(doc,
        ["Port", "Direction", "Width", "Description"],
        [
            ("in_pre", "Input", "1", "Right boundary input (tie to VDD = 1 for UP path)"),
            ("in_post", "Input", "1", "Left boundary input (tie to VSS = 0 for DOWN path)"),
            ("vsel<1:0>", "Input", "2", "Gray-coded shift control (see Section 3)"),
            ("sel_rst", "Input", "1", "Async reset enable (active high)"),
            ("vin_rst_b", "Input", "1", "Single-bit reset data (active-low) shared by all 4 muxes. "
             "vout = !vin_rst_b: assert low (0) → ON/1111, deassert high (1) → OFF/0000."),
            ("vout<3:0>", "Output", "4", "Thermometer code output"),
            ("vcca0p9", "Supply", "1", "0.9V supply"),
            ("vssx", "Supply", "1", "Ground"),
        ])

    # ══════════════════════════════════════════════════════════════════════
    # 2. ARCHITECTURE
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("2. Architecture", level=1)

    doc.add_heading("2.1 Mux Cell: hyp_adc_cdr_dac_mux_wr2", level=2)
    doc.add_paragraph(
        "Each mux_wr2 contains two cascaded inverting mux stages:"
    )
    add_mono(doc,
        "    data inputs ──→ [I_mux] ──voutb──→ [I_mux_rst] ──→ vout\n"
        "                       ↑                     ↑\n"
        "                      sel                sel_rst, vin_rst_b\n"
        "\n"
        "    I_mux:     voutb = !(sel ? in1 : in0)        ← data mux (inverting)\n"
        "    I_mux_rst: vout  = !(sel_rst ? vin_rst_b : voutb)  ← reset mux (inverting)\n"
        "\n"
        "    Normal mode (sel_rst=0):  vout = sel ? in1 : in0  (double inversion cancels)\n"
        "    Reset mode  (sel_rst=1):  vout = !vin_rst_b         (INVERTED reset!)\n"
        "\n"
        "    NOTE: All 4 muxes in the SR unit share the SAME vin_rst_b signal.\n"
        "    Therefore reset always produces all-0 or all-1 outputs.")

    doc.add_heading("2.2 Top-Level Connectivity", level=2)
    doc.add_paragraph(
        "The 4 muxes are wired in an asymmetric pattern that creates sliding "
        "cross-coupled pairs. The connections from the Spectre netlist are:"
    )
    add_table(doc,
        ["Instance", "in0 (sel=0)", "in1 (sel=1)", "sel", "Output"],
        [
            ("I_mux0", "in_pre", "vout1", "vsel<1>", "vout0"),
            ("I_mux1", "vout0", "vout2", "vsel<0>", "vout1"),
            ("I_mux2", "vout3", "vout1", "vsel<1>", "vout2"),
            ("I_mux3", "in_post", "vout2", "vsel<0>", "vout3"),
        ])

    doc.add_paragraph(
        "Key wiring observations:"
    )
    doc.add_paragraph(
        "• Even muxes (0, 2) share sel = vsel<1>; odd muxes (1, 3) share sel = vsel<0>",
        style="List Bullet"
    )
    doc.add_paragraph(
        "• in0 port of each mux — connections chain: in_pre → vout0 → vout3 → in_post "
        "(note mux2.in0 = vout3, not vout1 — this is the critical asymmetry)",
        style="List Bullet"
    )
    doc.add_paragraph(
        "• in1 port of each mux — connections cross-link: vout1 → vout2 → vout1 → vout2 "
        "(mux0 and mux2 both read vout1; mux1 and mux3 both read vout2)",
        style="List Bullet"
    )

    doc.add_heading("2.3 Block Diagram", level=2)
    block_diag_buf = generate_block_diagram()
    doc.add_picture(block_diag_buf, width=Inches(6.5))
    last_p = doc.paragraphs[-1]
    last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_mono(doc,
        "    in1 cross-links:  mux0.in1 = vout1,  mux1.in1 = vout2\n"
        "                      mux2.in1 = vout1,  mux3.in1 = vout2\n"
        "\n"
        "    in0 chain:        mux0.in0 = in_pre,   mux1.in0 = vout0\n"
        "                      mux2.in0 = vout3,    mux3.in0 = in_post", size=7)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 3. GRAY CODE OPERATION
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("3. Gray Code Operation", level=1)

    doc.add_heading("3.0 Bit Ordering Convention", level=2)
    doc.add_paragraph(
        "Throughout this document, the 4-bit thermometer output is written as "
        "[vout0, vout1, vout2, vout3]. In the physical layout, mux0 (vout0) is on the "
        "RIGHT and mux3 (vout3) is on the LEFT. The 'ON' region fills from the "
        "in_pre side (right/vout0) toward the in_post side (left/vout3):"
    )
    add_mono(doc,
        "    Code 0:  0000  (all OFF)\n"
        "    Code 1:  1000  (vout0 ON)\n"
        "    Code 2:  1100  (vout0, vout1 ON)\n"
        "    Code 3:  1110  (vout0, vout1, vout2 ON)\n"
        "    Code 4:  1111  (all ON)\n"
        "\n"
        "    Bit position:  [vout0] [vout1] [vout2] [vout3]\n"
        "    Physical:       (right)                 (left)\n"
        "    Boundary:       in_pre ──── fills this way ───→ in_post\n"
        "                     (=1)                            (=0)")

    doc.add_heading("3.1 What is a Cross-Coupled Pair?", level=2)
    doc.add_paragraph(
        "A 'pair' is formed when two muxes point at each other's output, creating "
        "a feedback loop that acts as a 1-bit latch. No clock is needed — the pair "
        "holds its value indefinitely through positive feedback, just like a "
        "cross-coupled inverter (SRAM cell)."
    )
    doc.add_paragraph(
        "How it works: When sel=1, a mux selects its in1 port. The in1 ports are "
        "cross-linked between pairs of muxes (vout1↔vout2 via mux0/mux2, and "
        "vout2↔vout1 via mux1/mux3). When BOTH muxes in a pair have sel=1, "
        "each reads the other's output — forming a stable feedback loop."
    )
    add_mono(doc,
        "    Example — vsel=10 (sel<1>=1, sel<0>=0), pair at (mux0 ↔ mux1):\n"
        "\n"
        "      mux0: sel<1>=1 → selects in1 = vout1 → so vout0 = vout1\n"
        "      mux1: sel<0>=0 → selects in0 = vout0 → so vout1 = vout0\n"
        "                                                    ↕\n"
        "                          vout0 = vout1 = vout0 = ... (LATCH!)\n"
        "\n"
        "    If previously vout0=1, vout1=1 → pair holds '1'\n"
        "    If previously vout0=0, vout1=0 → pair holds '0'\n"
        "\n"
        "    The pair is BISTABLE: it can hold either value. The value it actually\n"
        "    holds depends on what was written into it by the previous gray code step.")
    doc.add_paragraph(
        "Why does this matter? The pair value determines the thermometer code. "
        "The non-pair muxes are either forced by boundaries (in_pre=1, in_post=0) "
        "or they 'follow' the pair value. So changing the pair's value (by moving it "
        "to capture a boundary) shifts the thermometer code by ±1."
    )

    doc.add_heading("3.2 Pair Formation per vsel State", level=2)
    doc.add_paragraph(
        "At each vsel value, a different pair forms (or none at all). The table below "
        "shows which pair is active, what gets forced by boundaries, and which codes "
        "are holdable:"
    )
    add_table(doc,
        ["vsel<1:0>", "Cross-coupled pair", "Forced pins", "Followers",
         "Holdable codes"],
        [
            ("00", "None (no pair)", "vout0=in_pre=1, vout3=in_post=0",
             "vout1=vout0, vout2=vout3",
             "Code 2 only (forced)"),
            ("01", "mux2 ↔ mux3", "vout0 = in_pre = 1",
             "vout1 = pair value",
             "Code 1 (pair=0) or Code 4 (pair=1)"),
            ("10", "mux0 ↔ mux1", "vout3 = in_post = 0",
             "vout2 = pair value",
             "Code 0 (pair=0) or Code 3 (pair=1)"),
            ("11", "mux1 ↔ mux2", "—",
             "vout0 = pair, vout3 = pair",
             "Code 0 (pair=0) or Code 4 (pair=1)"),
        ])

    doc.add_heading("3.3 Pair Diagrams per vsel State", level=2)
    pair_diag_buf = generate_pair_diagrams()
    doc.add_picture(pair_diag_buf, width=Inches(6.5))
    last_p = doc.paragraphs[-1]
    last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("3.4 How Gray Code Steps Produce ±1", level=2)
    doc.add_paragraph(
        "Each gray code transition toggles exactly 1 bit of vsel, which moves the "
        "cross-coupled pair by one position. The pair captures the boundary value "
        "from the side it moves toward, extending or shrinking the thermometer code "
        "by 1."
    )

    # CW/CCW state diagram
    gray_diag_buf = generate_gray_code_diagram()
    doc.add_picture(gray_diag_buf, width=Inches(3.5))
    last_p = doc.paragraphs[-1]
    last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("Clockwise (CW) gray code = UP (+1):", style="List Bullet")
    add_mono(doc,
        "    00 ──→ 10 ──→ 11 ──→ 01 ──→ 00\n"
        "   (c2)   (c3)   (c4)   (sat)  (back to anchor)\n"
        "\n"
        "    From code 0 at vsel=11:\n"
        "    11(0) → 01(1) → 00(2) → 10(3) → 11(4) → 01(4,sat)")

    doc.add_paragraph("Counter-clockwise (CCW) gray code = DOWN (−1):", style="List Bullet")
    add_mono(doc,
        "    00 ──→ 01 ──→ 11 ──→ 10 ──→ 00\n"
        "   (c2)   (c1)   (c0)   (sat)  (back to anchor)\n"
        "\n"
        "    From code 4 at vsel=11:\n"
        "    11(4) → 10(3) → 00(2) → 01(1) → 11(0) → 10(0,sat)")

    doc.add_heading("3.5 Code-vsel Mapping", level=2)
    doc.add_paragraph(
        "Each thermometer code is associated with a specific vsel state during "
        "normal operation:"
    )
    add_table(doc,
        ["Code", "Thermometer [vout0..3]", "Active vsel", "Pair", "Pair value"],
        [
            ("0", "0000", "10 or 11", "(0↔1) or (1↔2)", "0"),
            ("1", "1000", "01", "(2↔3)", "0"),
            ("2", "1100", "00", "none (forced)", "—"),
            ("3", "1110", "10", "(0↔1)", "1"),
            ("4", "1111", "01 or 11", "(2↔3) or (1↔2)", "1"),
        ])
    doc.add_paragraph(
        "Note: Thermometer bits listed as [vout0, vout1, vout2, vout3]. "
        "Physically: vout0 (mux0) is on the RIGHT, vout3 (mux3) is on the LEFT. "
        "The '1' bits fill from the in_pre/mux0 side. Code 2 = 1100 means "
        "vout0=1, vout1=1, vout2=0, vout3=0.",
        style="List Bullet"
    )

    doc.add_heading("3.6 Shift Example: Code 2 → Code 3 (CW step 00→10)", level=2)
    add_mono(doc,
        "  BEFORE (vsel=00):          AFTER (vsel=10):\n"
        "  No pair, forced            Pair forms at (0↔1)\n"
        "\n"
        "  mux0=1 (from in_pre)       mux0 ←→ mux1 PAIR\n"
        "  mux1=1 (follows mux0)      Both were 1 → latch 1\n"
        "  mux2=0 (follows mux3)      mux2 follows pair → 1 (was 0!)\n"
        "  mux3=0 (from in_post)      mux3 = in_post = 0\n"
        "\n"
        "  [1,1,0,0] = code 2         [1,1,1,0] = code 3  (+1 ✓)\n"
        "\n"
        "  WHY: The pair captured '1' from the in_pre boundary.\n"
        "  mux2, no longer driven by boundary, follows the pair → flips to 1.\n"
        "  The block of 1s grew by one bit.", size=7)

    doc.add_heading("3.7 Simulation Waveforms", level=2)
    doc.add_paragraph(
        "The following waveforms are generated by running the netlist-verified "
        "Python simulation model (therm_sr_sim_v1_2)."
    )

    doc.add_paragraph("Full sweep — UP (CW gray code) then DOWN (CCW):",
                      style="List Bullet")
    sweep_buf = generate_sweep_waveform()
    doc.add_picture(sweep_buf, width=Inches(6.0))
    last_p = doc.paragraphs[-1]
    last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("Zigzag — UP 3, DOWN 2, UP 1 (direction reversal):",
                      style="List Bullet")
    zigzag_buf = generate_zigzag_waveform()
    doc.add_picture(zigzag_buf, width=Inches(5.0))
    last_p = doc.paragraphs[-1]
    last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 4. VALID AND INVALID INPUTS
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("4. Valid and Invalid Inputs", level=1)

    doc.add_heading("4.1 Valid vsel Transitions (Gray Code Only)", level=2)
    doc.add_paragraph(
        "Only single-bit transitions of vsel are valid. The gray code adjacency is:"
    )
    add_mono(doc,
        "           00\n"
        "          ╱  ╲\n"
        "        01    10\n"
        "          ╲  ╱\n"
        "           11\n"
        "\n"
        "    Valid transitions (1 bit changes):\n"
        "    00 ↔ 01    (vsel<0> toggles)\n"
        "    00 ↔ 10    (vsel<1> toggles)\n"
        "    01 ↔ 11    (vsel<1> toggles)\n"
        "    10 ↔ 11    (vsel<0> toggles)")

    doc.add_heading("4.2 INVALID vsel Transitions", level=2)
    p = doc.add_paragraph()
    run = p.add_run("NEVER transition 2 bits simultaneously:")
    run.bold = True
    run.font.color.rgb = RGBColor(200, 0, 0)

    add_table(doc,
        ["Transition", "Bits changed", "Risk", "Severity"],
        [
            ("00 ↔ 11", "2 (both)", "Both mux selects change — intermediate state "
             "has undefined pair. Race condition.", "CRITICAL — corrupts code"),
            ("01 ↔ 10", "2 (both)", "Same — transient state is either 00 or 11, "
             "causing an unintended shift.", "CRITICAL — corrupts code"),
        ],
        header_bg="F4CCCC")

    doc.add_heading("4.3 Valid Boundary Inputs", level=2)
    add_table(doc,
        ["Input", "Valid value", "Reason", "If wrong"],
        [
            ("in_pre", "1 (VDD)", "Feeds '1' from right (mux0 side) for UP shift",
             "UP shifts inject 0 instead of 1 — inverted behavior"),
            ("in_post", "0 (VSS)", "Feeds '0' from left (mux3 side) for DOWN shift",
             "DOWN shifts inject 1 instead of 0 — inverted behavior"),
        ])
    doc.add_paragraph(
        "Note: in_pre=1 and in_post=0 must be stable DC. If these float or glitch, "
        "the thermometer code will be corrupted."
    )

    doc.add_heading("4.4 Post-Saturation Behavior (DANGER)", level=2)
    p = doc.add_paragraph()
    run = p.add_run("WARNING: ")
    run.bold = True
    run.font.color.rgb = RGBColor(200, 0, 0)
    p.add_run(
        "If the gray code counter continues stepping after saturation, "
        "the code jumps by ±2 when passing through vsel=00 (anchor state)."
    )
    add_table(doc,
        ["Scenario", "Current state", "Next CW step", "Result", "Problem"],
        [
            ("UP past max", "code 4 @ vsel=01", "vsel=00",
             "Code 2 (forced)", "−2 jump!"),
            ("DOWN past min", "code 0 @ vsel=10", "vsel=00",
             "Code 2 (forced)", "+2 jump!"),
        ],
        header_bg="F4CCCC")
    doc.add_paragraph(
        "The grey controller MUST detect saturation (code 0 or 4) and stop the "
        "gray code counter. Otherwise the DAC output will oscillate."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 5. RESET
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("5. Reset Procedure", level=1)

    doc.add_heading("5.1 Reset Inversion", level=2)
    p = doc.add_paragraph()
    run = p.add_run("IMPORTANT: ")
    run.bold = True
    run.font.color.rgb = RGBColor(200, 0, 0)
    p.add_run(
        "The reset input is INVERTED by the hardware: all vout<3:0> = !vin_rst_b. "
        "Since vin_rst_b is a single bit shared across all 4 muxes, the unit can "
        "only reset to all-0 (code 0) or all-1 (code 4)."
    )
    add_table(doc,
        ["vin_rst_b", "All vout", "Unit code", "Release vsel"],
        [
            ("1", "0000", "0 (fully OFF)", "10 or 11"),
            ("0", "1111", "4 (fully ON)", "01 or 11"),
        ])
    doc.add_paragraph(
        "Since both code 0 and code 4 are valid at vsel=11, the only release "
        "vsel that works for a mixed array of ON/OFF units is vsel=11."
    )

    doc.add_heading("5.2 Why Toggling vsel During Reset is NOT Required", level=2)
    doc.add_paragraph(
        "Because vin_rst_b is shared (all muxes see the same value), the reset "
        "always produces a uniform output (all 0s or all 1s). At vsel=11, all "
        "mux sel=1, so every mux reads its in1 cross-link — which is another "
        "vout with the same reset value. This means voutb naturally aligns:"
    )
    add_mono(doc,
        "  During reset (sel_rst=1) at vsel=11:\n"
        "    All vout = !vin_rst_b  (forced by reset mux)\n"
        "    All in1  = vout_crosslink = !vin_rst_b (all same value)\n"
        "    All voutb = !(in1) = vin_rst_b   ← matches vin_rst_b!\n"
        "\n"
        "  On release (sel_rst 1→0):\n"
        "    vout = !voutb = !vin_rst_b = target   ✓ clean, no glitch\n"
        "\n"
        "  Conclusion: With shared vin_rst_b and vsel=11, no toggling is needed.")
    doc.add_paragraph(
        "However, if the design were modified to have per-mux in_rst with "
        "different values (e.g., in_rst = 0011 for code 2), toggling would be "
        "required — see Warning in Section 11."
    )

    doc.add_heading("5.3 Correct Reset Sequence", level=2)
    add_mono(doc,
        "  +------------------------------------------------------------+\n"
        "  |  1. Set vsel = 11  (MUST be 11 before and during reset!)   |\n"
        "  |  2. Set vin_rst_b per unit (0 -> ON/1111, 1 -> OFF/0000)     |\n"
        "  |  3. Assert sel_rst = 1                                     |\n"
        "  |  4. Wait >= t_settle for voutb to align                    |\n"
        "  |  5. Deassert sel_rst = 0   (no toggling needed)            |\n"
        "  +------------------------------------------------------------+\n"
        "\n"
        "  vsel:     ▔▔▔▔▔▔▔▔▔▔▔▔▔▔▔▔ 11 ▔▔▔▔▔▔▔▔▔▔▔▔▔▔▔▔▔▔▔▔▔\n"
        "  vin_rst_b:  XXXX< 0 or 1, stable >XXXXXXXXXXXXXXXXXXXXXXXXX\n"
        "  sel_rst:  _________╱▔▔▔▔▔▔▔▔▔▔▔▔╲________________________\n"
        "  vout:     XXXXXXXXX< !vin_rst_b >───────────────────────────\n"
        "\n"
        "  ⚠ CRITICAL: vsel MUST be 11 when sel_rst is released.\n"
        "  Releasing at any other vsel (especially 00) will FORCE the\n"
        "  wrong code immediately (vsel=00 always forces code 2).")

    doc.add_paragraph()
    doc.add_paragraph(
        "Simulated reset sequence (unit reset to code 0, in_rst=1111, "
        "vsel=11 held constant, no toggling):",
        style="List Bullet"
    )
    rst_buf = generate_reset_waveform()
    doc.add_picture(rst_buf, width=Inches(6.0))
    last_p = doc.paragraphs[-1]
    last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("5.4 Invalid Reset Sequences", level=2)
    add_table(doc,
        ["Sequence", "Problem", "Result"],
        [
            ("Release at vsel ≠ 11",
             "Pair forms at position incompatible with all-same-value code",
             "Some units may latch wrong code"),
            ("vin_rst_b changes during sel_rst=1",
             "Transient mismatch between vout and voutb",
             "Possible glitch on release"),
            ("sel_rst released before voutb settles",
             "voutb hasn't reached steady state",
             "Output glitches briefly before pair resolves"),
        ],
        header_bg="F4CCCC")

    doc.add_heading("5.5 Release vsel Rule", level=2)
    doc.add_paragraph(
        "With shared vin_rst_b, the only valid release vsel is 11. This is because "
        "vsel=11 creates a pair at (1↔2), and since all bits are uniform "
        "(all 0 or all 1), the pair always latches correctly."
    )
    doc.add_paragraph(
        "No other vsel value is safe as a universal release across a mixed "
        "array of units with different vin_rst_b values."
    )

    # ──────────────────────────────────────────────────────────────────────
    doc.add_heading("5.6 Full Operational Sequence", level=2)
    doc.add_paragraph(
        "This section describes the complete lifecycle of the SR block from "
        "power-on through normal shifting. This is the primary reference for "
        "system integration."
    )

    add_mono(doc,
        "  +--------------------------------------------------------------------+\n"
        "  |                     FULL OPERATION SEQUENCE                         |\n"
        "  +--------------------------------------------------------------------+\n"
        "  |                                                                    |\n"
        "  |  PHASE 0: STATIC SETUP (before reset)                             |\n"
        "  |    - in_pre = 1 (VDD)       -- right boundary (mux0), always high |\n"
        "  |    - in_post = 0 (VSS)      -- left boundary (mux3), always low   |\n"
        "  |    - These NEVER change during operation                           |\n"
        "  |                                                                    |\n"
        "  |  PHASE 1: RESET                                                   |\n"
        "  |    - Set vsel = 11                                                 |\n"
        "  |    - Set vin_rst_b per unit (0 = ON/1111, 1 = OFF/0000)             |\n"
        "  |    - Assert sel_rst = 1                                            |\n"
        "  |    - Wait >= t_settle                                              |\n"
        "  |    - Deassert sel_rst = 0                                          |\n"
        "  |    - Block is now at target code, vsel = 11, ready to shift        |\n"
        "  |                                                                    |\n"
        "  |  PHASE 2: FIRST SHIFT (from vsel=11)                              |\n"
        "  |    - To go UP:   toggle vsel<0> -> vsel = 01 (CW step)            |\n"
        "  |    - To go DOWN: toggle vsel<1> -> vsel = 10 (CCW step)           |\n"
        "  |                                                                    |\n"
        "  |  PHASE 3: NORMAL OPERATION (gray code stepping)                   |\n"
        "  |    - CW  (UP):  11 -> 01 -> 00 -> 10 -> 11 -> ... (+1 per step)  |\n"
        "  |    - CCW (DOWN): 11 -> 10 -> 00 -> 01 -> 11 -> ... (-1 per step) |\n"
        "  |    - Direction reversal: just reverse the gray code sequence       |\n"
        "  |    - STOP at saturation (code 0 or code 4 per unit)               |\n"
        "  |                                                                    |\n"
        "  +--------------------------------------------------------------------+", size=7)

    doc.add_heading("5.6.1 State After Reset", level=3)
    doc.add_paragraph(
        "Immediately after releasing sel_rst, the system is in this state:"
    )
    add_table(doc,
        ["Signal", "Value", "Notes"],
        [
            ("vsel<1:0>", "11", "Pair at (1↔2), both code 0 and code 4 are stable"),
            ("sel_rst", "0", "Normal operation mode"),
            ("vin_rst_b", "don't care", "Ignored when sel_rst=0"),
            ("in_pre", "1", "Constant — right boundary (mux0 side)"),
            ("in_post", "0", "Constant — left boundary (mux3 side)"),
            ("vout<3:0>", "0000 or 1111", "Per-unit, set by vin_rst_b during reset"),
        ])

    doc.add_heading("5.6.2 First Gray Code Step", level=3)
    doc.add_paragraph(
        "The first shift always starts from vsel=11. The direction determines "
        "which bit to toggle:"
    )
    add_table(doc,
        ["Direction", "Toggle", "New vsel", "Effect on boundary unit"],
        [
            ("UP (CW)", "vsel<0>: 1→0", "10",
             "Pair moves to (0↔1), captures '1' from in_pre → +1"),
            ("DOWN (CCW)", "vsel<1>: 1→0", "01",
             "Pair moves to (2↔3), captures '0' from in_post → −1"),
        ])
    doc.add_paragraph(
        "Subsequent steps continue the gray code sequence. The controller "
        "must track the current vsel to know which bit to toggle next."
    )

    doc.add_heading("5.6.3 Saturation Handling", level=3)
    doc.add_paragraph(
        "When a unit reaches code 0 (all OFF) or code 4 (all ON), further "
        "steps in the same direction are absorbed — the code clips. "
        "The controller MUST detect saturation and stop the gray counter, "
        "because one additional step past saturation passes through vsel=00 "
        "(the anchor state) and causes a ±2 code jump."
    )
    add_mono(doc,
        "  Normal saturation (safe):\n"
        "    code 4 @ vsel=11 → CW → vsel=01: code stays 4 (clipped)\n"
        "    code 0 @ vsel=11 → CCW → vsel=10: code stays 0 (clipped)\n"
        "\n"
        "  DANGER — one more step past saturation:\n"
        "    code 4 @ vsel=01 → CW → vsel=00: code JUMPS to 2 (−2!)\n"
        "    code 0 @ vsel=10 → CCW → vsel=00: code JUMPS to 2 (+2!)\n"
        "\n"
        "  The gray controller must stop before vsel reaches 00 at extremes.", size=7)

    doc.add_heading("5.6.4 Direction Reversal", level=3)
    doc.add_paragraph(
        "Reversing direction is immediate — just reverse the gray code rotation. "
        "No special setup is needed. The cross-coupled pair correctly captures "
        "the new boundary value on the very next step."
    )
    add_mono(doc,
        "  Example: UP 2, then DOWN 1 (starting from code 2 @ vsel=00)\n"
        "\n"
        "    Step 1: CW  00→10 → code 3\n"
        "    Step 2: CW  10→11 → code 4\n"
        "    Step 3: CCW 11→10 → code 3   (immediate reversal, no dead step)\n"
        "\n"
        "  The net result is +1, as expected (2 + 2 − 1 = 3).", size=7)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 6. COMPLETE TRUTH TABLES
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("6. Complete Truth Tables", level=1)

    doc.add_heading("6.1 State Transition Table (Normal Operation)", level=2)
    doc.add_paragraph(
        "For each (current code, current vsel) → single gray step → (next code, next vsel):"
    )
    add_table(doc,
        ["From code", "From vsel", "CW step\n(UP)", "New code", "CCW step\n(DOWN)", "New code"],
        [
            ("0", "11", "→ 01", "1  (+1)", "→ 10", "0  (sat)"),
            ("0", "10", "→ 11", "0  (sat)", "→ 00", "2  (+2 DANGER)"),
            ("1", "01", "→ 00", "2  (+1)", "→ 11", "0  (−1)"),
            ("2", "00", "→ 10", "3  (+1)", "→ 01", "1  (−1)"),
            ("3", "10", "→ 11", "4  (+1)", "→ 00", "2  (−1)"),
            ("4", "11", "→ 01", "4  (sat)", "→ 10", "3  (−1)"),
            ("4", "01", "→ 00", "2  (−2 DANGER)", "→ 11", "4  (sat)"),
        ])

    doc.add_paragraph(
        "Gray shaded rows show saturation and the dangerous ±2 jump if the "
        "controller continues past saturation."
    )

    doc.add_heading("6.2 Reset Truth Table", level=2)
    doc.add_paragraph(
        "With shared vin_rst_b, only code 0 and code 4 are achievable per unit. "
        "Release vsel must be 11 for both:")
    add_table(doc,
        ["Target", "vin_rst_b", "Release\nvsel=00", "Release\nvsel=01",
         "Release\nvsel=10", "Release\nvsel=11"],
        [
            ("code 0 (OFF)", "1", "FAIL→c2", "FAIL→c1", "PASS ✓", "PASS ✓"),
            ("code 4 (ON)", "0", "FAIL→c2", "PASS ✓", "FAIL→c3", "PASS ✓"),
        ])
    doc.add_paragraph(
        "Both rows pass at vsel=11 → use vsel=11 as the universal release state. "
        "No vsel toggling is required (see Section 7.5)."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 7. SCALING TO FULL DAC
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("7. Scaling to 9-Bit DAC (511 codes)", level=1)
    doc.add_paragraph(
        "Each 4-mux SR unit handles 5 thermometer codes (0 through 4). "
        "For a 9-bit DAC with 511 unit current sources, approximately 128 SR units "
        "are chained together:"
    )
    add_mono(doc,
        "  Unit 0           Unit 1           Unit 2           Unit 127\n"
        "  +------+         +------+         +------+         +------+\n"
        "  |      |--vout3->|      |--vout3->|      | . . .  |      |\n"
        "  |  SR  | in_post |  SR  | in_post |  SR  |        |  SR  |-> in_post=0\n"
        "  |      |<-in_pre-|      |<-in_pre-|      |        |      |\n"
        "  +------+         +------+         +------+         +------+\n"
        "  in_pre=1\n"
        "\n"
        "  All units share the same vsel<1:0> and sel_rst.\n"
        "  Each unit has its own vin_rst_b (1-bit per unit).\n"
        "  Total codes: 128 x 4 + 1 = 513 >= 511 needed.")

    doc.add_paragraph(
        "The chaining works because each unit's boundary outputs (vout0/vout3) "
        "connect to the adjacent unit's boundary inputs (in_pre/in_post). The "
        "thermometer boundary slides across unit boundaries seamlessly."
    )

    doc.add_heading("7.1 Inter-Unit Port Mapping", level=2)
    doc.add_paragraph(
        "Adjacent SR units connect through their boundary ports. "
        "The 1→0 thermometer boundary propagates across unit boundaries "
        "identically to intra-unit propagation:"
    )
    add_table(doc,
        ["Source", "Signal", "Destination", "Purpose"],
        [
            ("Unit[n].vout3", "therm_out", "Unit[n+1].in_post",
             "Passes '1' or '0' into next unit (left side)"),
            ("Unit[n+1].vout0", "therm_out", "Unit[n].in_pre",
             "Passes '1' or '0' into previous unit (right side)"),
            ("Global", "vsel<1:0>", "All Unit[*].vsel",
             "Same gray code control to all units"),
            ("Global", "sel_rst", "All Unit[*].sel_rst",
             "Shared reset enable"),
            ("Per-unit", "vin_rst_b", "Each Unit[n].vin_rst_b",
             "1-bit per unit: 0 (→ON/1111) or 1 (→OFF/0000)"),
        ])

    doc.add_heading("7.2 Boundary Crossing Example", level=2)
    doc.add_paragraph(
        "When the thermometer boundary sits at the edge of a unit, "
        "one gray code step moves it into the adjacent unit:"
    )
    add_mono(doc,
        "  Before (code = 4 in Unit[n], code = 0 in Unit[n+1]):\n"
        "     Unit[n]:   [1, 1, 1, 1]    Unit[n+1]: [0, 0, 0, 0]\n"
        "     vout3=1 ─────▶ in_post=1   (Unit[n+1] sees '1' at its left boundary)\n"
        "\n"
        "  After CW step (vsel toggles):\n"
        "     Unit[n]:   [1, 1, 1, 1]    Unit[n+1]: [1, 0, 0, 0]\n"
        "     The '1' from Unit[n].vout3 propagates into Unit[n+1].vout3,\n"
        "     then chains through mux3→mux2→... exactly as within a single unit.\n"
        "\n"
        "  Effective: global thermometer code increased by 1 (unit[n] stayed full,\n"
        "  unit[n+1] gained one '1')", size=7)

    doc.add_heading("7.3 Reset in Multi-Unit Configuration", level=2)
    doc.add_paragraph(
        "Each unit has a single vin_rst_b input shared across all 4 internal muxes. "
        "All units share the same sel_rst and vsel<1:0>. "
        "The reset code decoder sets each unit's vin_rst_b based on whether the "
        "unit falls below or above the thermometer boundary:"
    )
    doc.add_paragraph(
        "Units below the boundary (fully ON): vin_rst_b = 0 → vout = 1111",
        style="List Bullet"
    )
    doc.add_paragraph(
        "Units above the boundary (fully OFF): vin_rst_b = 1 → vout = 0000",
        style="List Bullet"
    )
    doc.add_paragraph(
        "Constraint: the global reset code must be a multiple of 4 (0, 4, 8, ..., 508, 512). "
        "Partial thermometer codes within a unit are not achievable with shared vin_rst_b."
    )

    doc.add_heading("7.4 Reset to Mid-Code 256 (Design Target)", level=2)
    doc.add_paragraph(
        "The CDR DAC resets to mid-code 256 (out of 511). Since 256 is exactly "
        "divisible by 4, every SR unit resets to either all-1s or all-0s — "
        "no unit holds a partial thermometer code."
    )
    add_table(doc,
        ["Units", "Count", "vin_rst_b", "vout (= !vin_rst_b)", "State"],
        [
            ("Unit 0 – 63", "64", "0", "1111", "Fully ON"),
            ("Unit 64 – 127", "64", "1", "0000", "Fully OFF"),
        ])
    doc.add_paragraph(
        "Total thermometer code = 64 × 4 = 256 ones. "
        "This corresponds to exactly half-scale output."
    )
    doc.add_paragraph(
        "Since all units are either code 0 (0000) or code 4 (1111), "
        "the release vsel must be valid for both states simultaneously:"
    )
    add_table(doc,
        ["Unit state", "Valid release vsel", "Reason"],
        [
            ("0000 (code 0)", "10 or 11",
             "Pair at (0↔1) or (1↔2), all bits = 0"),
            ("1111 (code 4)", "01 or 11",
             "Pair at (2↔3) or (1↔2), all bits = 1"),
            ("Intersection", "11 only",
             "Only vsel=11 satisfies both"),
        ],
        header_bg="C6EFCE")
    p = doc.add_paragraph()
    run = p.add_run("Release vsel MUST be 11 (vsel<1:0> = 11) for mid-code 256 reset.")
    run.bold = True

    doc.add_heading("7.5 No Toggling Required", level=2)
    doc.add_paragraph(
        "Because vin_rst_b is shared across all 4 muxes, every unit always resets "
        "to a uniform state (all 0s or all 1s). The voutb nodes inside each mux "
        "naturally align with vin_rst_b at vsel=11 — no vsel toggling is needed."
    )
    doc.add_paragraph(
        "At vsel=11, all mux sel=1, so every mux selects its in1 (cross-link) "
        "input. When all 4 vout bits are identical:"
    )
    add_mono(doc,
        "  voutb = !(sel=1 ? in1 : in0) = !(vout_crosslink) = !vout = vin_rst_b\n"
        "\n"
        "  Unit at code 4 (vout=1111, vin_rst_b=0):\n"
        "    All in1 = 1 → all voutb = 0 = vin_rst_b   ✓ (4/4 match)\n"
        "\n"
        "  Unit at code 0 (vout=0000, vin_rst_b=1):\n"
        "    All in1 = 0 → all voutb = 1 = vin_rst_b   ✓ (4/4 match)\n"
        "\n"
        "  On release: vout = !voutb = !vin_rst_b = target   ✓ clean, no glitch", size=7)

    p = doc.add_paragraph()
    run = p.add_run(
        "Since vin_rst_b is a single shared signal, every unit is always "
        "fully uniform during reset. This guarantee eliminates the need for "
        "vsel toggling entirely — voutb always aligns with vin_rst_b at vsel=11."
    )
    run.font.size = Pt(9)
    run.italic = True

    doc.add_heading("7.6 Mid-Code 256 Reset Sequence", level=2)
    add_mono(doc,
        "  +------------------------------------------------------------+\n"
        "  |  1. Set vsel = 11                                          |\n"
        "  |  2. Set vin_rst_b per unit:                                  |\n"
        "  |       Units 0-63:   vin_rst_b = 0  (-> vout = 1111)         |\n"
        "  |       Units 64-127: vin_rst_b = 1  (-> vout = 0000)         |\n"
        "  |  3. Assert sel_rst = 1                                     |\n"
        "  |  4. Wait for voutb to settle (>= 1 x t_cell)              |\n"
        "  |  5. Deassert sel_rst = 0       (no toggling needed)        |\n"
        "  +------------------------------------------------------------+\n"
        "\n"
        "  vsel:     ▔▔▔▔▔▔▔▔▔▔▔▔▔▔▔▔ 11 ▔▔▔▔▔▔▔▔▔▔▔▔▔▔▔▔▔▔▔▔\n"
        "  vin_rst_b:  XXXX< per-unit 0 or 1 (stable) >XXXXXXX\n"
        "  sel_rst:  _________╱▔▔▔▔▔▔▔▔▔▔▔╲____________________\n"
        "  vout:     XXXXXXXXX< 256 ones | 256 zeros >──────────", size=7)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 8. RISKS AND MITIGATIONS
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("8. Risks and Mitigations", level=1)

    add_table(doc,
        ["#", "Risk", "Cause", "Impact", "Mitigation"],
        [
            ("1", "Post-saturation ±2 jump",
             "Gray counter continues past code 0 or 4",
             "DAC output jumps 2 codes",
             "Grey controller must detect saturation and stop"),
            ("2", "2-bit vsel transition",
             "Both vsel bits change simultaneously",
             "Undefined intermediate state, code corruption",
             "Ensure gray-code sequencing in controller"),
            ("3", "Reset value lost on release",
             "sel_rst released at vsel ≠ 11",
             "Wrong initial DAC code → CDR starts at wrong freq",
             "Always release at vsel=11"),
            ("4", "Reset polarity",
             "vin_rst_b sense confused (vout = !vin_rst_b)",
             "DAC resets to inverted code",
             "vin_rst_b=0 → ON (1111), vin_rst_b=1 → OFF (0000)"),
            ("5", "Metastability in cross-coupled pair",
             "Pair forms from muxes with different values",
             "Output may take long to resolve or resolve wrong",
             "Only occurs at design errors; normal gray steps "
             "always form pairs from agreeing muxes"),
            ("6", "Boundary input glitch",
             "in_pre or in_post not stable",
             "Wrong value injected at thermometer boundary",
             "Tie in_pre=VDD, in_post=VSS with decoupling"),
            ("7", "Non-thermometer state",
             "SEU, noise, or illegal reset",
             "DAC output undefined, may not recover via shifting",
             "Reset after anomaly detection"),
        ])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 9. TIMING CONSTRAINTS
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("9. Timing Constraints", level=1)

    doc.add_heading("9.1 Mux Propagation Delay", level=2)
    doc.add_paragraph(
        "Each mux_wr2 cell contains two cascaded inverting muxes "
        "(transmission gate + output inverter). The propagation delay through "
        "one cell is:"
    )
    add_mono(doc,
        "  t_mux = t_TG + t_INV  (per stage)\n"
        "  t_cell = 2 × t_mux    (two cascaded stages: I_mux + I_mux_rst)\n"
        "\n"
        "  At TSMC 2nm, SLVT, 0.9V, TT corner:\n"
        "    t_TG  ≈ 5–10 ps  (transmission gate)\n"
        "    t_INV ≈ 5–8 ps   (inverter)\n"
        "    t_cell ≈ 20–36 ps per mux_wr2 cell")

    doc.add_heading("9.2 Cross-Coupled Pair Settling", level=2)
    doc.add_paragraph(
        "When a gray code step moves the cross-coupled pair position, "
        "the new pair must resolve through positive feedback. The settling time "
        "depends on the regeneration time constant of the cross-coupled inverters:"
    )
    add_mono(doc,
        "  t_settle = t_cell × N_loops\n"
        "\n"
        "  Worst case: pair resolves from near-metastable state\n"
        "    t_regen ≈ τ × ln(V_DD / ΔV_initial)\n"
        "\n"
        "  Typical: pair captures a strong '1' or '0' from boundary\n"
        "    t_settle ≈ 2–3 × t_cell ≈ 40–108 ps\n"
        "\n"
        "  Note: In normal gray code operation, the pair always captures\n"
        "  a rail-to-rail boundary value, so metastability does NOT occur.\n"
        "  Settling is fast (2–3 loop delays).")

    doc.add_heading("9.3 Multi-Unit Propagation", level=2)
    doc.add_paragraph(
        "In the chained 128-unit configuration, a single gray code step only "
        "changes the thermometer code by ±1. The change propagates through at "
        "most 1 unit boundary (2 mux cells: vout3 → in_post → mux3 → mux2...). "
        "There is no ripple carry across multiple units."
    )
    add_mono(doc,
        "  Worst-case single-step delay (boundary crossing):\n"
        "    t_step = t_cell (source unit) + t_wire + t_cell (dest unit)\n"
        "           ≈ 2 × t_cell + t_wire\n"
        "           ≈ 40–72 ps + routing delay")

    doc.add_heading("9.4 vsel Setup / Hold Constraints", level=2)
    doc.add_paragraph(
        "The vsel signals are the select inputs of transmission-gate muxes. "
        "They must be stable when the mux output is being sampled by the "
        "cross-coupled pair."
    )
    add_table(doc,
        ["Parameter", "Constraint", "Notes"],
        [
            ("t_setup(vsel)", "> t_settle",
             "vsel must be stable before pair resolves"),
            ("t_hold(vsel)", "> 0",
             "vsel must remain stable until pair is latched"),
            ("t_vsel_skew", "< t_cell / 2",
             "Max skew between vsel<1> and vsel<0> "
             "(to prevent transient 2-bit change)"),
            ("Min vsel period", "> t_settle + margin",
             "Allow pair to resolve fully before next toggle"),
        ])

    doc.add_heading("9.5 Reset Timing", level=2)
    add_table(doc,
        ["Parameter", "Constraint", "Notes"],
        [
            ("t_rst_setup(vin_rst_b)", "> t_cell before sel_rst↑",
             "vin_rst_b must be stable before asserting reset"),
            ("t_rst_hold(vin_rst_b)", "> 0 after sel_rst↓",
             "vin_rst_b must hold through release"),
            ("t_rst_min_width", "> t_settle",
             "Reset must stay asserted long enough "
             "for voutb to settle"),
            ("t_release_setup(vsel)", "> t_cell before sel_rst↓",
             "vsel=11 must be stable before releasing reset"),
        ])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 10. SPECTRE NETLIST REFERENCE
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("10. Spectre Netlist Reference", level=1)
    doc.add_paragraph(
        "The following Spectre subcircuit definitions are the authoritative "
        "source for the block connectivity. All analysis in this document was "
        "verified against these netlists."
    )

    doc.add_heading("10.1 Bare Mux: hyp_adc_cdr_dac_mux2b_unit", level=2)
    doc.add_paragraph(
        "Transmission-gate 2:1 mux with output inverter. "
        "Function: voutb = !(sel ? in1 : in0). "
        "6 ports (selb generated internally)."
    )
    add_mono(doc,
        "  subckt hyp_adc_cdr_dac_mux2b_unit in0 in1 sel vcca0p9 voutb vssx\n"
        "\n"
        "  // Transmission gate mux (TG pass gates + output inverter)\n"
        "  // sel=0: passes in0 through TG (PMOS gate=sel, NMOS gate=selb)\n"
        "  // sel=1: passes in1 through TG (PMOS gate=selb, NMOS gate=sel)\n"
        "  // Internal inverter drives voutb\n"
        "\n"
        "  // Net function: voutb = !(sel ? in1 : in0)\n"
        "  // Note: selb is generated internally (no external selb port)", size=7)

    doc.add_heading("10.2 Mux Wrapper: hyp_adc_cdr_dac_mux_wr2", level=2)
    doc.add_paragraph(
        "Two cascaded inverting muxes: data mux → reset mux. "
        "The double inversion cancels in normal mode."
    )
    add_mono(doc,
        "  subckt hyp_adc_cdr_dac_mux_wr2 in0 in1 in_rst sel sel_rst vcca0p9 vout vssx\n"
        "\n"
        "  I_mux (in0 in1 sel vcca0p9 voutb vssx)\n"
        "      hyp_adc_cdr_dac_mux2b_unit\n"
        "  // voutb = !(sel ? in1 : in0)\n"
        "\n"
        "  I_mux_rst (voutb in_rst sel_rst vcca0p9 vout vssx)\n"
        "      hyp_adc_cdr_dac_mux2b_unit\n"
        "  // vout = !(sel_rst ? in_rst : voutb)\n"
        "\n"
        "  // Normal (sel_rst=0): vout = !(voutb) = !!(sel?in1:in0) = sel?in1:in0\n"
        "  // Reset  (sel_rst=1): vout = !(in_rst) = !in_rst", size=7)

    doc.add_heading("10.3 SR Unit: hyp_adc_cdr_dac_sr_unit", level=2)
    doc.add_paragraph(
        "4 mux_wr2 instances with cross-coupled feedback wiring. "
        "All 4 muxes share a single vin_rst_b input."
    )
    add_mono(doc,
        "  // From Spectre netlist (hyp_adc_cdr_dac_sr_unit schematic):\n"
        "\n"
        "  I_mux0 (in_pre  vout1 vin_rst_b vsel<1> sel_rst vcca0p9 vout0 vssx)\n"
        "      hyp_adc_cdr_dac_mux_wr2\n"
        "  //  in0=in_pre, in1=vout1, in_rst=vin_rst_b, sel=vsel<1> → vout0\n"
        "\n"
        "  I_mux1 (vout0 vout2 vin_rst_b vsel<0> sel_rst vcca0p9 vout1 vssx)\n"
        "      hyp_adc_cdr_dac_mux_wr2\n"
        "  //  in0=vout0, in1=vout2, in_rst=vin_rst_b, sel=vsel<0> → vout1\n"
        "\n"
        "  I_mux2 (vout3 vout1 vin_rst_b vsel<1> sel_rst vcca0p9 vout2 vssx)\n"
        "      hyp_adc_cdr_dac_mux_wr2\n"
        "  //  in0=vout3, in1=vout1, in_rst=vin_rst_b, sel=vsel<1> → vout2\n"
        "  //  NOTE: in0=vout3 (not vout1) — asymmetric wiring!\n"
        "\n"
        "  I_mux3 (in_post vout2 vin_rst_b vsel<0> sel_rst vcca0p9 vout3 vssx)\n"
        "      hyp_adc_cdr_dac_mux_wr2\n"
        "  //  in0=in_post, in1=vout2, in_rst=vin_rst_b, sel=vsel<0> → vout3\n"
        "\n"
        "  // KEY: All 4 muxes share the SAME vin_rst_b signal.\n"
        "  // This guarantees uniform reset (all 0 or all 1).", size=7)

    doc.add_heading("10.4 Pin Ordering Convention", level=2)
    doc.add_paragraph(
        "The hyp_adc_cdr_dac_mux_wr2 subcircuit port order is:"
    )
    add_mono(doc,
        "  (in0 in1 in_rst sel sel_rst vcca0p9 vout vssx)\n"
        "   ─┬─ ─┬─ ──┬── ─┬─ ──┬──── ───┬──── ─┬── ─┬──\n"
        "    │   │    │    │    │         │      │    │\n"
        "    │   │    │    │    │         │      │    └ Ground\n"
        "    │   │    │    │    │         │      └ Output\n"
        "    │   │    │    │    │         └ Supply (0.9V)\n"
        "    │   │    │    │    └ Reset mux select\n"
        "    │   │    │    └ Data mux select\n"
        "    │   │    └ Reset data input (shared as vin_rst_b at SR level)\n"
        "    │   └ Data input 1 (selected when sel=1)\n"
        "    └ Data input 0 (selected when sel=0)", size=7)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 11. WARNING: SPLIT in_rst (NON-UNIFORM RESET)
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("11. Warning: Split in_rst (Non-Uniform Reset)", level=1)

    p = doc.add_paragraph()
    run = p.add_run("⚠ WARNING — APPLICABILITY LIMITATION")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(200, 0, 0)

    doc.add_paragraph(
        "This document assumes the production design where all 4 muxes in each "
        "SR unit share a single vin_rst_b signal. If a future design modification "
        "splits in_rst into separate per-mux signals (in_rst<3:0>) with "
        "different values, the following sections of this document are NO LONGER "
        "VALID without additional analysis:"
    )
    doc.add_paragraph(
        "Section 5.2 (no toggling required) — would need toggling",
        style="List Bullet"
    )
    doc.add_paragraph(
        "Section 5.3 (simple reset sequence) — would need full gray cycle toggle",
        style="List Bullet"
    )
    doc.add_paragraph(
        "Section 5.5 (vsel=11 as universal release) — may not be valid for all codes",
        style="List Bullet"
    )
    doc.add_paragraph(
        "Section 7.5 (no toggling proof) — proof relies on uniform vout",
        style="List Bullet"
    )
    doc.add_paragraph(
        "Section 7.3 (reset code must be multiple of 4) — partial codes become possible",
        style="List Bullet"
    )

    doc.add_heading("11.1 Example: Split in_rst = 0011 (target code 2)", level=2)
    doc.add_paragraph(
        "If in_rst were split per-mux to achieve code 2 (vout = 1100), "
        "the reset input would be in_rst<3:0> = 0011 (inverted). Here's why "
        "the simple vsel=11 reset FAILS:"
    )
    add_mono(doc,
        "  Target: vout = [1, 1, 0, 0] = code 2\n"
        "  Split in_rst: mux0.in_rst=0, mux1.in_rst=0, mux2.in_rst=1, mux3.in_rst=1\n"
        "\n"
        "  During reset at vsel=11 (all sel=1, selecting in1 cross-links):\n"
        "    vout = [1, 1, 0, 0]  (forced by per-mux in_rst: !0,!0,!1,!1)\n"
        "\n"
        "  But voutb depends on in1 = vout_crosslink:\n"
        "    mux0: in1 = vout1 = 1 → voutb0 = !(1) = 0   need in_rst[0]=0  ✓\n"
        "    mux1: in1 = vout2 = 0 → voutb1 = !(0) = 1   need in_rst[1]=0  ✗ MISMATCH\n"
        "    mux2: in1 = vout1 = 1 → voutb2 = !(1) = 0   need in_rst[2]=1  ✗ MISMATCH\n"
        "    mux3: in1 = vout2 = 0 → voutb3 = !(0) = 1   need in_rst[3]=1  ✓\n"
        "\n"
        "  Result: 2 out of 4 voutb nodes DON'T match in_rst!\n"
        "  On release: vout = !voutb = [1, 0, 1, 0] ≠ [1, 1, 0, 0]  ← WRONG!\n"
        "\n"
        "  FIX: Must toggle vsel through a full gray cycle during reset\n"
        "  to force all voutb nodes to settle correctly, then release at\n"
        "  the correct hold vsel for code 2 (which is vsel=00, NOT 11).", size=7)

    doc.add_paragraph(
        "This example demonstrates why the shared vin_rst_b architecture was chosen: "
        "it eliminates the voutb-alignment problem entirely by guaranteeing "
        "uniform outputs during reset."
    )

    # ══════════════════════════════════════════════════════════════════════
    # SAVE
    # ══════════════════════════════════════════════════════════════════════
    output_path = os.path.join(os.path.dirname(__file__),
                               "Thermometer_SR_Block_Description.docx")
    doc.save(output_path)
    print(f"Document generated: {output_path}")
    return output_path


if __name__ == "__main__":
    try:
        generate()
        print("Done.")
    except Exception as e:
        print(f"ERROR: {e}")
        raise

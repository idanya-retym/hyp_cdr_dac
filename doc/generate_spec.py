"""
Generate DAC CDR Specification Document (.docx)

Creates a full spec template for the HYP Gen1 CDR DAC with:
- Architecture, Electrical, Interface, Timing, and Test sections
- Retym header placeholder and revision table

Usage:
    python generate_spec.py
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
import os
from datetime import date


def set_cell_text(cell, text, bold=False, size=10):
    """Helper to set cell text with formatting."""
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold
    run.font.size = Pt(size)


def add_table_with_headers(doc, headers, rows=None, col_widths=None):
    """Add a formatted table with headers and optional data rows."""
    n_cols = len(headers)
    n_rows = 1 + (len(rows) if rows else 3)  # header + data/placeholder rows
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=10)
        shading = table.rows[0].cells[i]._element
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        tc_pr = shading.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "D9E2F3")
        shd.set(qn("w:val"), "clear")
        tc_pr.append(shd)

    # Data rows
    if rows:
        for r_idx, row_data in enumerate(rows):
            for c_idx, val in enumerate(row_data):
                set_cell_text(table.rows[r_idx + 1].cells[c_idx], str(val), size=10)
    else:
        # Leave placeholder rows empty
        for r_idx in range(1, n_rows):
            for c_idx in range(n_cols):
                set_cell_text(table.rows[r_idx].cells[c_idx], "", size=10)

    doc.add_paragraph()  # spacing
    return table


def generate_spec():
    doc = Document()

    # -- Page setup --
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # =========================================================================
    # TITLE PAGE
    # =========================================================================
    doc.add_paragraph()  # spacing

    # Company header placeholder
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[RETYM LOGO]")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Retym, Inc.")
    run.font.size = Pt(16)
    run.bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("HYP Gen1 CDR DAC\nSpecification Document")
    run.font.size = Pt(24)
    run.bold = True

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("9-Bit Fully Thermometric Differential Current-Steering DAC")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # Document info table
    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = "Table Grid"
    info_data = [
        ("Document Number", "RETYM-HYP-DAC-SPEC-001"),
        ("Version", "0.1 (Draft)"),
        ("Date", date.today().strftime("%Y-%m-%d")),
        ("Author", "[Author Name]"),
        ("Status", "DRAFT"),
    ]
    for i, (label, value) in enumerate(info_data):
        set_cell_text(info_table.rows[i].cells[0], label, bold=True, size=11)
        set_cell_text(info_table.rows[i].cells[1], value, size=11)

    doc.add_page_break()

    # =========================================================================
    # REVISION HISTORY
    # =========================================================================
    doc.add_heading("Revision History", level=1)

    rev_headers = ["Rev", "Date", "Author", "Description"]
    rev_rows = [
        ("0.1", "2025-07-01", "[Author]", "Initial draft – skeleton/template"),
        ("0.2", date.today().strftime("%Y-%m-%d"), "[Author]", "Updated to 9-bit (511 codes); added INL/DNL formulas (sec 4.2)"),
    ]
    add_table_with_headers(doc, rev_headers, rev_rows)

    doc.add_page_break()

    # =========================================================================
    # TABLE OF CONTENTS placeholder
    # =========================================================================
    doc.add_heading("Table of Contents", level=1)
    p = doc.add_paragraph()
    p.add_run("[Update field to generate TOC in Word: Ctrl+A → F9]").font.color.rgb = RGBColor(128, 128, 128)
    doc.add_page_break()

    # =========================================================================
    # 1. INTRODUCTION
    # =========================================================================
    doc.add_heading("1. Introduction", level=1)

    doc.add_heading("1.1 Purpose", level=2)
    doc.add_paragraph(
        "This document defines the specification for the Digital-to-Analog Converter (DAC) "
        "used in the HYP Gen1 Clock and Data Recovery (CDR) system. The DAC converts the "
        "digital loop filter output into an analog control voltage/current for the VCO/CCO."
    )

    doc.add_heading("1.2 Scope", level=2)
    doc.add_paragraph(
        "This spec covers the DAC architecture, electrical performance targets, "
        "digital interface definition, timing requirements, and test/characterization plan."
    )

    doc.add_heading("1.3 References", level=2)
    ref_headers = ["Ref #", "Document", "Description"]
    ref_rows = [
        ("1", "[CDR System Spec]", "HYP Gen1 CDR top-level system specification"),
        ("2", "[Process PDK]", "Foundry design manual and device models"),
        ("3", "[Layout Guidelines]", "Current source matching layout rules"),
    ]
    add_table_with_headers(doc, ref_headers, ref_rows)

    doc.add_heading("1.4 Definitions & Abbreviations", level=2)
    abbr_headers = ["Term", "Definition"]
    abbr_rows = [
        ("CDR", "Clock and Data Recovery"),
        ("DAC", "Digital-to-Analog Converter"),
        ("DNL", "Differential Non-Linearity"),
        ("INL", "Integral Non-Linearity"),
        ("LSB", "Least Significant Bit"),
        ("VCO", "Voltage Controlled Oscillator"),
        ("CCO", "Current Controlled Oscillator"),
        ("BBPD", "Bang-Bang Phase Detector"),
        ("FLL", "Frequency-Locked Loop"),
    ]
    add_table_with_headers(doc, abbr_headers, abbr_rows)

    doc.add_page_break()

    # =========================================================================
    # 2. SYSTEM OVERVIEW
    # =========================================================================
    doc.add_heading("2. System Overview", level=1)

    doc.add_heading("2.1 CDR Architecture Context", level=2)
    doc.add_paragraph(
        "[Block diagram showing CDR loop: PD → Digital Loop Filter → DAC → VCO/CCO → ...]"
    )
    doc.add_paragraph(
        "The DAC resides between the digital loop filter and the analog VCO/CCO. "
        "It must translate digital frequency/phase correction commands into a precise "
        "analog control signal with minimal latency and noise."
    )

    doc.add_heading("2.2 DAC Role in CDR Loop", level=2)
    doc.add_paragraph("[Describe the DAC's role in loop dynamics, bandwidth, and jitter.]")

    doc.add_heading("2.3 Top-Level Block Diagram", level=2)
    doc.add_paragraph("[Insert DAC top-level block diagram here]")

    doc.add_page_break()

    # =========================================================================
    # 3. ARCHITECTURE
    # =========================================================================
    doc.add_heading("3. Architecture", level=1)

    doc.add_heading("3.1 DAC Topology", level=2)
    doc.add_paragraph(
        "Topology: 9-bit fully thermometric, differential current-steering DAC.\n"
        "Number of unit elements: N = 2^9 - 1 = 511 unit current sources."
    )

    doc.add_heading("3.2 Unit Current Source Design", level=2)
    doc.add_paragraph("[Describe unit cell: transistor type, dimensions, cascode, etc.]")

    doc.add_heading("3.3 Switching Architecture", level=2)
    doc.add_paragraph("[Describe current-steering switch pair, return-to-zero options, etc.]")

    doc.add_heading("3.4 Bias Generation", level=2)
    doc.add_paragraph("[Describe bias current reference, distribution, and filtering.]")

    doc.add_heading("3.5 Output Stage", level=2)
    doc.add_paragraph("[Describe output termination, load, compliance voltage range.]")

    doc.add_page_break()

    # =========================================================================
    # 4. ELECTRICAL SPECIFICATIONS
    # =========================================================================
    doc.add_heading("4. Electrical Specifications", level=1)

    doc.add_heading("4.1 Key Performance Parameters", level=2)

    spec_headers = ["Parameter", "Symbol", "Min", "Typ", "Max", "Unit", "Notes"]
    spec_rows = [
        ("Resolution", "N", "—", "9", "—", "bits", "Thermometric"),
        ("Number of codes", "—", "—", "511", "—", "—", "2^N - 1"),
        ("Full-scale current", "I_FS", "", "10.02", "", "mA", "511 × 19.6 µA"),
        ("LSB current", "I_LSB", "", "19.6", "", "µA", "I_FS / 511"),
        ("DNL", "DNL", "", "", "±0.5", "LSB", "Target yield TBD"),
        ("INL", "INL", "", "", "", "LSB", ""),
        ("Current source mismatch", "σ(Iu)/Iu", "", "", "", "%", "See matching analysis"),
        ("Output compliance", "V_comp", "", "", "", "V", ""),
        ("Supply voltage", "VDD", "", "", "", "V", ""),
        ("Power consumption", "P_DAC", "", "", "", "mW", ""),
        ("Settling time", "t_settle", "", "", "", "ps", "To 0.5 LSB"),
        ("Update rate", "f_update", "", "", "", "GHz", ""),
        ("Output noise (integrated)", "I_n", "", "", "", "pA_rms", "BW TBD"),
        ("PSRR", "PSRR", "", "", "", "dB", "At freq TBD"),
        ("Temperature range", "T", "", "", "", "°C", ""),
    ]
    add_table_with_headers(doc, spec_headers, spec_rows)

    # ----- 4.2 INL and DNL Definitions -----
    doc.add_heading("4.2 INL and DNL Definitions", level=2)

    doc.add_paragraph(
        "For a fully thermometric differential current-steering DAC with N = 2^B - 1 "
        "unit current sources (B = 9, N = 511):"
    )

    doc.add_heading("4.2.1 DNL (Differential Non-Linearity)", level=3)
    doc.add_paragraph(
        "DNL measures the deviation of each code step from the ideal 1-LSB step size.\n\n"
        "Definition (endpoint-fitted):\n\n"
        "    DNL(k) = [V_out(k+1) - V_out(k)] / LSB_actual  -  1\n\n"
        "where:\n"
        "    LSB_actual = [V_out(N) - V_out(0)] / N\n\n"
        "For a fully thermometric DAC, each code transition activates exactly one unit "
        "current source. Therefore:\n\n"
        "    DNL(k) = [I_unit(k) - I_avg] / I_avg\n\n"
        "where I_unit(k) is the current of the k-th unit source and I_avg is the mean "
        "of all N unit sources.\n\n"
        "Specification: |DNL| < 0.5 LSB (guarantees monotonicity)"
    )

    doc.add_heading("4.2.2 INL (Integral Non-Linearity)", level=3)
    doc.add_paragraph(
        "INL measures the cumulative deviation of the transfer function from the ideal "
        "straight line.\n\n"
        "Definition (endpoint-fitted):\n\n"
        "    INL(k) = [V_out(k) - V_out(0)] / LSB_actual  -  k\n\n"
        "Equivalently, INL is the cumulative sum of DNL:\n\n"
        "    INL(k) = Σ_{i=0}^{k-1} DNL(i)\n\n"
        "For a differential DAC with output V_diff = I_diff × R_L:\n\n"
        "    INL(k) = [V_diff(k) - V_diff(0)] / LSB_actual  -  k\n"
        "    LSB_actual = [V_diff(N) - V_diff(0)] / N"
    )

    doc.add_heading("4.2.3 Differential DAC: Finite Rout Effect", level=3)
    doc.add_paragraph(
        "For a differential current-steering DAC, the first-order finite-Rout effect "
        "cancels between the two output branches, producing only a gain error (no INL):\n\n"
        "    Gain Error = (N × R_L) / (2 × R_out)\n\n"
        "Second-order INL from Rout variation across code:\n\n"
        "    INL_diff_max ≈ (N² / 42) × Δg_ds × R_L   [LSBs]\n\n"
        "where Δg_ds is the variation of output conductance across the output voltage "
        "swing."
    )

    doc.add_heading("4.2.4 Mismatch-Limited DNL (Statistical)", level=3)
    doc.add_paragraph(
        "Given relative mismatch σ_rel = σ(I_unit)/I_unit, the probability that ALL "
        "N codes meet |DNL| < DNL_target:\n\n"
        "    Yield = [2·Φ(DNL_target / σ_rel) - 1]^N\n\n"
        "where Φ is the standard normal CDF.\n\n"
        "Solving for required σ_rel at a given yield target Y:\n\n"
        "    σ_rel = DNL_target / Φ⁻¹[(1 + Y^(1/N)) / 2]\n\n"
        "For N = 511, |DNL| < 0.5 LSB:\n"
        "  • 90% yield → σ_rel ≤ TBD %\n"
        "  • 95% yield → σ_rel ≤ TBD %\n"
        "  • 99% yield → σ_rel ≤ TBD %\n"
        "  • 3σ (99.7%) yield → σ_rel ≤ TBD %"
    )

    # ----- 4.3 Matching Requirements -----
    doc.add_heading("4.3 Matching Requirements", level=2)
    doc.add_paragraph(
        "Based on Monte Carlo analysis (see dac_sigma_vs_dnl.py):\n\n"
        "For |DNL| < 0.5 LSB with 511 unit elements:\n"
        "  • 90% yield → σ_rel ≤ TBD %\n"
        "  • 95% yield → σ_rel ≤ TBD %\n"
        "  • 99% yield → σ_rel ≤ TBD %\n"
        "  • 99.7% yield → σ_rel ≤ TBD %"
    )

    doc.add_heading("4.3 Noise Budget", level=2)
    doc.add_paragraph("[Define noise contributors: thermal, flicker, supply, reference.]")

    doc.add_heading("4.4 Power Budget", level=2)
    doc.add_paragraph("[Breakdown of power by sub-block.]")

    doc.add_page_break()

    # =========================================================================
    # 5. DIGITAL INTERFACE
    # =========================================================================
    doc.add_heading("5. Digital Interface", level=1)

    doc.add_heading("5.1 Interface Type", level=2)
    doc.add_paragraph(
        "[Define interface: async pulse (UP/DN), synchronous parallel, or hybrid.\n"
        "Reference system_questions.md Q1–Q9 for open decisions.]\n\n"
        "CODE[8:0]: 9-bit thermometer state (511 codes)"
    )

    doc.add_heading("5.2 Signal List", level=2)
    sig_headers = ["Signal", "Direction", "Width", "Description"]
    sig_rows = [
        ("UP", "Input", "1", "Increment DAC code by 1 LSB"),
        ("DN", "Input", "1", "Decrement DAC code by 1 LSB"),
        ("RESET_N", "Input", "1", "Active-low async reset to initial code"),
        ("CODE[8:0]", "Output/Internal", "9", "Current DAC thermometer code (optional readback)"),
        ("SAT_HI", "Output", "1", "Saturation flag – code at maximum"),
        ("SAT_LO", "Output", "1", "Saturation flag – code at minimum"),
    ]
    add_table_with_headers(doc, sig_headers, sig_rows)

    doc.add_heading("5.3 Protocol & Timing Diagram", level=2)
    doc.add_paragraph("[Insert timing diagram for UP/DN pulse interface.]")
    doc.add_paragraph(
        "• Minimum pulse width: TBD ps\n"
        "• Minimum pulse spacing: TBD ps\n"
        "• Setup/hold to internal clock (if sync): TBD ps"
    )

    doc.add_heading("5.4 Reset & Power-Up Behavior", level=2)
    doc.add_paragraph(
        "[Define initial DAC code after reset: mid-code (64), min (0), or programmable.\n"
        "Reference system_questions.md Q4.]"
    )

    doc.add_heading("5.5 Saturation Handling", level=2)
    doc.add_paragraph(
        "[Define behavior at code 0 and 127: clip, flag, or both.\n"
        "Reference system_questions.md Q5.]"
    )

    doc.add_heading("5.6 Clock Domain & Synchronization", level=2)
    doc.add_paragraph(
        "[Define clock domain relationship between digital loop filter and DAC.\n"
        "Reference system_questions.md Q8.]"
    )

    doc.add_page_break()

    # =========================================================================
    # 6. TIMING SPECIFICATIONS
    # =========================================================================
    doc.add_heading("6. Timing Specifications", level=1)

    doc.add_heading("6.1 Latency Budget", level=2)
    timing_headers = ["Path", "Latency", "Unit", "Notes"]
    timing_rows = [
        ("UP/DN pulse to current change", "", "ps", "Interface + decoder + switch"),
        ("Reset to mid-code settled", "", "ns", ""),
        ("Code-to-code settling (1 LSB)", "", "ps", ""),
        ("Full-scale settling (0 → 127)", "", "ns", ""),
    ]
    add_table_with_headers(doc, timing_headers, timing_rows)

    doc.add_heading("6.2 Update Rate", level=2)
    doc.add_paragraph("[Maximum pulse rate / clock frequency for code updates.]")

    doc.add_heading("6.3 Glitch Energy", level=2)
    doc.add_paragraph("[Max allowable glitch at code transitions, in LSB·ps or pV·s.]")

    doc.add_page_break()

    # =========================================================================
    # 7. LAYOUT & MATCHING
    # =========================================================================
    doc.add_heading("7. Layout & Matching Considerations", level=1)

    doc.add_heading("7.1 Unit Cell Layout", level=2)
    doc.add_paragraph("[Common-centroid, interdigitated, or other arrangement.]")

    doc.add_heading("7.2 Dummy Elements", level=2)
    doc.add_paragraph("[Number and placement of dummy devices at array edges.]")

    doc.add_heading("7.3 Routing & Shielding", level=2)
    doc.add_paragraph("[Signal routing strategy, ground shielding, guard rings.]")

    doc.add_heading("7.4 Area Estimate", level=2)
    doc.add_paragraph("[Estimated die area: TBD µm × TBD µm.]")

    doc.add_page_break()

    # =========================================================================
    # 8. TEST & CHARACTERIZATION
    # =========================================================================
    doc.add_heading("8. Test & Characterization Plan", level=1)

    doc.add_heading("8.1 Production Tests", level=2)
    test_headers = ["Test", "Method", "Spec Limit", "Notes"]
    test_rows = [
        ("DNL", "Code-by-code ramp", "< 0.5 LSB", ""),
        ("INL", "Cumulative DNL", "TBD", ""),
        ("Full-scale current", "Force code 511, measure I_out", "TBD", ""),
        ("Power consumption", "Measure supply current", "TBD", ""),
        ("Monotonicity", "Verify no code reversal (DNL > -1)", "Pass/Fail", ""),
    ]
    add_table_with_headers(doc, test_headers, test_rows)

    doc.add_heading("8.2 Characterization Tests", level=2)
    char_headers = ["Test", "Conditions", "Purpose"]
    char_rows = [
        ("DNL vs. temperature", "–40°C to 125°C", "Verify matching over temp"),
        ("DNL vs. supply", "VDD ± 10%", "PSRR impact on linearity"),
        ("Settling time", "1 LSB and full-scale step", "Verify timing spec"),
        ("Noise spectral density", "Mid-code, measure output", "Verify noise budget"),
        ("Glitch energy", "Mid-code transition", "Verify glitch spec"),
    ]
    add_table_with_headers(doc, char_headers, char_rows)

    doc.add_heading("8.3 Bench Measurement Setup", level=2)
    doc.add_paragraph("[Describe test equipment, PCB requirements, measurement methodology.]")

    doc.add_page_break()

    # =========================================================================
    # 9. OPEN ITEMS & TBDs
    # =========================================================================
    doc.add_heading("9. Open Items & TBDs", level=1)

    doc.add_paragraph(
        "The following items require resolution (see doc/system_questions.md):"
    )

    open_headers = ["Item #", "Question", "Owner", "Status"]
    open_rows = [
        ("1", "Step size: ±1 LSB only or variable?", "", "Open"),
        ("2", "Frequency acquisition mode (FLL) exists?", "", "Open"),
        ("3", "Max UP/DN pulse rate?", "", "Open"),
        ("4", "Initial DAC code on reset?", "", "Open"),
        ("5", "Saturation behavior (clip/flag/both)?", "", "Open"),
        ("6", "UP/DN mutual exclusivity guaranteed?", "", "Open"),
        ("7", "Code readback needed?", "", "Open"),
        ("8", "Clock domain relationship?", "", "Open"),
        ("9", "Supply / IO voltage levels?", "", "Open"),
    ]
    add_table_with_headers(doc, open_headers, open_rows)

    # =========================================================================
    # SAVE
    # =========================================================================
    output_path = os.path.join(os.path.dirname(__file__), "HYP_Gen1_CDR_DAC_Spec.docx")
    doc.save(output_path)
    print(f"Spec document generated: {output_path}")
    return output_path


if __name__ == "__main__":
    try:
        path = generate_spec()
        print("Done.")
    except Exception as e:
        print(f"ERROR: {e}")
        raise
